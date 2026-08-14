"""
工具层：定义莲心AI可调用的所有工具（OpenAI / DeepSeek Function Calling 格式）
每个工具包含两部分：
  1. TOOL_DEFINITIONS  — 发送给 DeepSeek API 的工具描述（OpenAI格式）
  2. 对应的 Python 执行函数
"""

import json
import re
import subprocess
import threading
import time
from pathlib import Path
from datetime import datetime
from typing import Optional
import sys
import logging
import fnmatch
import difflib
import os
import litellm
from brain.code_intel import goto_definition, find_references, get_diagnostics

# ── 子代理可用工具白名单（第二阶段） ──────────────────────
# delegate_task 生成的子代理只能使用这些工具
_SUBAGENT_ALLOWED_TOOLS = {
    "read_file", "read_file_lines", "read_file_chunk",
    "grep_file", "search_code", "glob_files",
    "search_files_everything", "list_directory", "edit_file", "code_structure",
    "run_shell", "run_python_code", "diff_files", "git_status",
}

logger = logging.getLogger("brain.tools")

from utils.paths import get_user_data_dir
from brain.document_cache import MarkdownDocumentCache

# 记忆系统（统一使用 SQLite 后端）
from brain.graph_memory import (
    # 分类事实 CRUD（替换 long_term.json）
    add_fact as _memory_add,
    add_memory_fragment as _memory_add_fragment,
    search_facts as _memory_search,
    update_facts as _memory_update,
    delete_facts as _memory_delete,
    list_all_facts,
    migrate_from_json,
    # 统一搜索（事实 + 图边）
    unified_search,
    format_unified_search_result,
    # 五元组图查询
    search_graph_ranked,
    query_by_entity,
    query_connected,
    format_graph_result,
    get_graph_stats,
    delete_entity,
    get_fact_by_id,
    get_fact_fragments,
)
# 格式化工具和常量仍然从 memory_store 取（无存储依赖）
from brain.graph_memory import format_all_memories
from brain.current_state import (
    set_current_state as _state_set,
    update_current_state as _state_update,
    resolve_current_state as _state_resolve,
    list_current_states as _state_list,
)
# 每块最大字符数（read_file 默认读第0块，read_file_chunk 可读任意块）
_CHUNK_SIZE = 15000
_document_cache: MarkdownDocumentCache | None = None
_document_cache_lock = threading.Lock()

# QQ 桥接 worker 引用（由 main_window 启动时注册）
_qq_bridge_worker = None

def _register_qq_bridge(worker):
    """注册 QQBridgeWorker 实例，供 send_file_to_qq 工具使用。"""
    global _qq_bridge_worker
    _qq_bridge_worker = worker

# 肩载摄像头（ESP32-CAM WebSocket bridge，由 GUI 启动后注册）
_shoulder_bridge = None

def _register_shoulder_bridge(bridge):
    """注册 HardwareBridge 实例，供摄像头/云台工具使用。"""
    global _shoulder_bridge
    _shoulder_bridge = bridge

def _get_shoulder_bridge():
    """获取肩载摄像头桥接实例。首次调用时自动创建。"""
    global _shoulder_bridge
    if _shoulder_bridge is None:
        from brain.hardware_bridge import HardwareBridge
        _shoulder_bridge = HardwareBridge()
    return _shoulder_bridge

# 文本类扩展名（用编码检测读取）
_TEXT_EXTENSIONS = {
    ".txt", ".md", ".py", ".js", ".ts", ".html", ".css",
    ".json", ".yaml", ".yml", ".xml", ".csv", ".ini",
    ".cfg", ".toml", ".log", ".bat", ".sh", ".c", ".cpp",
    ".h", ".java", ".rs", ".go", ".rb", ".php",
}

# 待办管理器（稍后初始化）
_todo_manager = None
_music_info_callback = None
_music_control_callback = None
_note_refresh_callback = None
_proactive_toggle_callback = None  # 主动聊天开关变更后通知调度器刷新
_expression_callback = None  # Galgame 模式：切换立绘表情

# 跨端搜索上下文（thread-local，由 AgentCore 在调用 execute_tool 前设置）
_tool_context = threading.local()

# 日记消息源（由 GUI/QQ 桥接在调用工具前设置，提供当日对话消息）
_diary_message_source = None  # Callable[[], List[Dict]] 返回 [{"role": ..., "content": ...}, ...]

# ── 代理配置工具函数 ──────────────────────────────────────────
def _get_proxies() -> dict | None:
    """读取用户代理配置，返回 requests 兼容的 proxies 字典，未启用或未配置返回 None。"""
    try:
        from config import get_proxy_config
        cfg = get_proxy_config()
        if not cfg.get("enabled", False):
            return None
        http = cfg.get("http_proxy", "").strip()
        https = cfg.get("https_proxy", "").strip()
        proxies = {}
        if http:
            proxies["http"] = http
        if https:
            proxies["https"] = https
        if not proxies:
            return None
        no_proxy = cfg.get("no_proxy", "").strip()
        if no_proxy:
            proxies["no_proxy"] = no_proxy
        return proxies or None
    except Exception:
        return None



def set_cross_session_context(session_id: int, history_mgr, model: str = ""):
    """设置当前线程的跨端搜索上下文（供 search_cross_session 工具使用）。"""
    _tool_context.cross_session = {
        "session_id": session_id,
        "history_mgr": history_mgr,
        "model": str(model or ""),
    }


def _current_memory_provenance() -> dict:
    """Capture the current persisted user turn for memory writes and corrections."""
    provenance = {
        "source_session_id": None,
        "source_channel": "",
        "source_message_ids": [],
        "persona_id": "",
        "occurred_at": "",
        "review_model": "",
    }
    ctx = getattr(_tool_context, "cross_session", None)
    if ctx is not None:
        provenance["source_session_id"] = ctx.get("session_id")
        provenance["review_model"] = ctx.get("model", "")
        try:
            session = ctx["history_mgr"].get_session(provenance["source_session_id"])
            provenance["source_channel"] = (session or {}).get("channel", "")
            recent_messages = ctx["history_mgr"].get_messages_with_ids(
                provenance["source_session_id"], limit=4
            )
            for message in reversed(recent_messages):
                if message.get("role") == "user":
                    provenance["source_message_ids"] = [int(message["id"])]
                    provenance["occurred_at"] = message.get("timestamp", "")
                    break
        except Exception:
            pass
    try:
        from brain.persona.runtime import capture_persona_snapshot
        provenance["persona_id"] = capture_persona_snapshot().profile.id
    except Exception:
        pass
    return provenance


# ── DeepSeek/OpenAI 工具定义 ────────────────────────────────
TOOL_DEFINITIONS = [
    {
        "type": "function",
        "function": {
            "name": "read_diary",
            "description": (
                "读取莲心时间胶囊中的真实日记和共同书页。用户提到昨天、前天、某天日记或时间胶囊时必须优先使用，"
                "不得用 read_file 或文件搜索代替；返回内容会标注数据库来源、日期和作者。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "date": {"type": "string", "description": "YYYY-MM-DD，也可传‘昨天’或‘前天’"},
                    "keyword": {"type": "string", "description": "在时间胶囊日记中检索的关键词"},
                    "limit": {"type": "integer", "description": "最多返回几篇，默认1"},
                },
                "additionalProperties": False,
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_diary",
            "description": "让莲心根据今天的共同记录写入时间胶囊日记；已有日记时不会静默覆盖。",
            "parameters": {
                "type": "object",
                "properties": {
                    "message_count": {"type": "integer"},
                    "force": {"type": "boolean"},
                },
                "additionalProperties": False,
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "discover_connections",
            "description": (
                "图谱关系发现引擎：从指定实体出发，遍历知识图谱，发现所有直接和间接关联的实体、关系、路径。"
                "与 search_graph_memory 不同——discover 是图遍历（发现你不知道的关联），search 是关键词搜索（找你知道的）。"
                "适用于「我和XX有什么关系」「这个项目涉及哪些人」「把所有相关的东西找出来」。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "entity_name": {
                        "type": "string",
                        "description": "起始实体名称，如'用户'、'莲心AI'、'张三'。默认'用户'"
                    },
                    "depth": {
                        "type": "integer",
                        "description": "遍历深度（1=直接关联, 2=两跳间接关联, 最多3），默认2"
                    }
                },
                "required": ["entity_name"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": (
                "读取指定路径文件的内容（第0块，即开头部分）。"
                "⚠️ 此工具只接受文件路径，不接受目录路径。如需浏览目录中的文件列表，请先用 list_directory 或 search_files_everything。"
                "支持 .txt .md .py .csv .json 等文本文件（自动识别 UTF-8/GBK 等编码），"
                "以及 .docx Word文档、.pdf PDF文件。"
                "每次最多返回 15000 字符。若文件更长，返回结果中会注明总块数，"
                "此时应告知用户文件较长并询问是否继续阅读，如需继续则调用 read_file_chunk。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "文件的绝对路径或相对路径，例如 C:/Users/user/Desktop/note.txt"
                    }
                },
                "required": ["path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_file_chunk",
            "description": (
                "读取长文件的指定分块。当 read_file 提示文件有多块时，"
                "使用此工具读取后续内容。chunk_index 从 0 开始，"
                "0 表示开头，1 表示第二块，以此类推。"
                "适用于小说、长报告等篇幅较大的文档。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "文件路径，与 read_file 中的路径相同"
                    },
                    "chunk_index": {
                        "type": "integer",
                        "description": "要读取的块编号，从 0 开始"
                    }
                },
                "required": ["path", "chunk_index"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": "将指定内容写入文件。文件不存在则创建，已存在则覆盖。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "目标文件的路径"
                    },
                    "content": {
                        "type": "string",
                        "description": "要写入文件的文本内容"
                    }
                },
                "required": ["path", "content"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "list_directory",
            "description": "列出指定目录下的所有文件和子目录。找不到文件时启用 recursive=True 递归搜索子目录。找文件先 list_directory 确认，再打开查看。\n找到文件后如需查看内容，请立即调用 read_file，不要只报告文件名。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "目录路径，不填则默认列出桌面"
                    },
                    "recursive": {
                        "type": "boolean",
                        "description": "是否递归列出子目录，默认 False，找不到文件时设为 True"
                    },
                    "max_depth": {
                        "type": "integer",
                        "description": "递归深度限制，默认 3"
                    }
                },
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "search_files_everything",
            "description": (
                "【推荐优先使用】毫秒级全盘文件搜索（依赖 Everything + es.exe）。"
                "支持按扩展名、最近修改天数过滤。若 Everything 未索引目标目录，自动降级为 Python 直接搜索。"
                "找到文件后如需查看内容，请立即调用 read_file，不要只报告文件名。"
                "调用示例：search_files_everything(keyword='简历', ext='docx;pdf', recent_days=2)"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "keyword": {
                        "type": "string",
                        "description": "文件名关键词，如 '简历'"
                    },
                    "ext": {
                        "type": "string",
                        "description": "扩展名过滤，分号分隔，如 'pdf;docx'。不填则不过滤"
                    },
                    "folder": {
                        "type": "string",
                        "description": "限定搜索目录，如 'C:\\Users\\me\\Desktop'。不填则全盘搜索"
                    },
                    "recent_days": {
                        "type": "integer",
                        "description": "限定最近 N 天修改的文件，如 1 表示今天，2 表示昨天至今。0 表示不限制"
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "最多返回条数，默认 20"
                    }
                },
                "required": ["keyword"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_file_info_everything",
            "description": "获取文件的元数据：大小、修改时间、创建时间。",
            "parameters": {
                "type": "object",
                "properties": {
                    "filepath": {
                        "type": "string",
                        "description": "文件的完整路径"
                    }
                },
                "required": ["filepath"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "save_memory",
            "description": (
                "将重要信息永久保存到长期记忆，下次启动程序后仍然记得。"
                "触发时机：① 用户明确说'记住这个'/'帮我记下来'等；"
                "② 用户透露姓名、职业、重要项目、明显的个人偏好时可主动记录。"
                "每条记忆用简洁的一句话描述，例如：'用户的名字叫小明'。"
                "记忆会自动归入合适的分类：profile(个人档案)、preferences(偏好)、events(事件)、"
                "knowledge(知识)、behaviors(行为模式)、skills(技能)——你可以指定或用默认。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "fact": {
                        "type": "string",
                        "description": "要记住的事实，一句话，简洁明了"
                    },
                    "category": {
                        "type": "string",
                        "enum": ["profile", "preferences", "events", "knowledge", "behaviors", "skills"],
                        "description": "记忆分类，不填则自动归为 knowledge"
                    }
                },
                "required": ["fact"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "open_app",
            "description": (
                "打开指定的应用程序、文件夹或文件。"
                "支持常见应用别名（记事本、计算器、画图、资源管理器、命令行、任务管理器等），"
                "也支持直接传入完整路径（如 C:/Program Files/xxx/xxx.exe 或文件夹路径）。"
                "用户说'打开微信'、'帮我启动记事本'、'打开桌面'等时调用此工具。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {
                        "type": "string",
                        "description": "应用名称（如'记事本'、'微信'）或完整路径（如 C:/Windows/notepad.exe）"
                    }
                },
                "required": ["name"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_clipboard",
            "description": (
                "读取用户当前剪贴板中的文字内容并返回。"
                "当用户说'帮我看看我复制的内容'、'分析一下剪贴板里的东西'、"
                "'我刚复制了一段代码/文字，帮我…'等时调用此工具。"
            ),
            "parameters": {
                "type": "object",
                "properties": {},
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "run_command",
            "description": "在系统中执行一个 shell 命令。仅支持白名单中的安全命令。",
            "parameters": {
                "type": "object",
                "properties": {
                    "command": {
                        "type": "string",
                        "description": "要执行的命令，例如 'dir C:/Users/user/Desktop'"
                    }
                },
                "required": ["command"]
            }
        }
    },
    # ==================== 时间工具 ====================
    {
        "type": "function",
        "function": {
            "name": "get_current_time",
            "description": "获取当前的日期、时间、星期、农历、节假日等信息。当用户询问现在几点、今天几号、今天星期几、农历日期、今天是不是节假日等问题时调用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "format": {
                        "type": "string",
                        "description": "返回格式，可选：full(完整信息)、date(仅日期)、time(仅时间)、weekday(仅星期)、lunar(仅农历)、holiday(仅节假日)",
                        "enum": ["full", "date", "time", "weekday", "lunar", "holiday"]
                    }
                },
                "required": []
            }
        }
    },
    # ==================== 余额查询工具 ====================
    {
        "type": "function",
        "function": {
            "name": "get_balance",
            "description": (
                "查询 DeepSeek API 账户的当前余额。"
                "当用户询问余额、还剩多少钱、账户余额、还有多少余额、余额够不够等问题时调用此工具。"
                "该工具会实时从 DeepSeek 官方接口获取最新余额信息。"
            ),
            "parameters": {
                "type": "object",
                "properties": {},
                "required": []
            }
        }
    },
    # ==================== 待办清单工具 ====================
    {
        "type": "function",
        "function": {
            "name": "add_todo",
            "description": (
                "【必须调用】当用户要求添加待办事项、设置提醒、记录任务时，必须使用此工具，绝对不要直接回复用户说'已添加'。"
                "用户话语示例：'提醒我明天下午3点开会'、'添加待办 买牛奶'、'帮我记一下 明天上午10点打电话'、'设置一个提醒 晚上8点吃药'。"
                "调用此工具后，系统会真实存储待办，然后你才能告知用户已添加。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {
                        "type": "string",
                        "description": "待办标题，应简洁明确，例如'开会'、'买牛奶'"
                    },
                    "due_time": {
                        "type": "string",
                        "description": "截止时间，ISO格式字符串（例如2026-04-21T15:00:00）。如果用户没有明确时间，则为null"
                    },
                    "priority": {
                        "type": "string",
                        "enum": ["high", "medium", "low"],
                        "description": "优先级，根据用户描述判断：高优先级（紧急、重要）、低优先级（不着急），否则为medium"
                    },
                    "description": {
                        "type": "string",
                        "description": "详细描述，可选"
                    }
                },
                "required": ["title"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "list_todos",
            "description": "列出当前所有未完成的待办事项。当用户问'有哪些待办'、'显示我的待办清单'、'我还有什么事要做'时调用。",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "complete_todo",
            "description": "将某个待办标记为完成。当用户说'完成...'、'标记...为完成'、'...做好了'时调用。如果用户提供的标题关键词匹配到多个待办，请询问用户选择哪一个。",
            "parameters": {
                "type": "object",
                "properties": {
                    "title_keyword": {
                        "type": "string",
                        "description": "待办标题中的关键词，用于模糊匹配"
                    }
                },
                "required": ["title_keyword"]
            }
        }
    },

    {
        "type": "function",
        "function": {
            "name": "read_excel",
            "description": "读取 Excel 文件（.xlsx）的内容，返回表格数据（文本格式）。支持指定工作表名称和最大行数。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Excel 文件的绝对路径"
                    },
                    "sheet_name": {
                        "type": "string",
                        "description": "工作表名称，默认第一个工作表"
                    },
                    "max_rows": {
                        "type": "integer",
                        "description": "最多读取的行数，默认 100"
                    }
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "run_python_code",
            "description": "在安全沙箱中执行 Python 代码，返回标准输出和错误。适用于计算、数据处理等任务。不支持文件操作或系统命令。",
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {
                        "type": "string",
                        "description": "要执行的 Python 代码"
                    },
                    "timeout": {
                        "type": "integer",
                        "description": "超时时间（秒），默认 10"
                    }
                },
                "required": ["code"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "write_excel",
            "description": "将数据写入 Excel 文件（.xlsx）。可创建新文件或覆盖已有文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "目标 Excel 文件的绝对路径"},
                    "sheet_name": {"type": "string", "description": "工作表名称，默认 'Sheet1'"},
                    "data": {"type": "array", "description": "二维数组，每行是一个列表，例如 [['姓名','年龄'],['张三',28]]"}
                },
                "required": ["file_path", "data"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "copy_excel_content",
            "description": "将源 Excel 文件的内容复制到目标 Excel 文件（支持指定工作表）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "source_path": {"type": "string", "description": "源文件路径"},
                    "target_path": {"type": "string", "description": "目标文件路径"},
                    "source_sheet": {"type": "string", "description": "源工作表名称，默认第一个"},
                    "target_sheet": {"type": "string", "description": "目标工作表名称，默认与源相同或新建"}
                },
                "required": ["source_path", "target_path"]
            }
        }
    },

    # ==================== Word 文档写入工具 ====================
    {
        "type": "function",
        "function": {
            "name": "write_docx",
            "description": (
                "将文本内容写入 Word 文档（.docx）。可创建新文件或覆盖已有文件。"
                "当用户要求「生成 Word 文档」、「创建 .docx 文件」、「把内容整理到 Word」时使用此工具。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "目标 .docx 文件的绝对路径，例如 C:/Users/用户名/Desktop/结果.docx"
                    },
                    "content": {
                        "type": "string",
                        "description": "要写入的文本内容，支持多段。若需要分段，请用两个换行符 \\n\\n 分隔。"
                    }
                },
                "required": ["file_path", "content"]
            }
        }
    },


    {
        "type": "function",
        "function": {
            "name": "format_document",
            "description": (
                "【推荐】将Markdown格式的文本内容转换为格式优美的Word文档（.docx）。"
                "当用户要求'生成报告'、'写文档'、'整理内容'、'排版'时，应优先使用此工具。"
                "你可以直接输出Markdown格式的内容，此工具会将其转换为专业排版的Word文档。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "content": {
                        "type": "string",
                        "description": "要转换的Markdown格式文本内容。"
                    },
                    "output_path": {
                        "type": "string",
                        "description": "要生成的Word文档的保存路径。"
                    },
                    "use_template": {
                        "type": "boolean",
                        "description": "是否使用预设的Word模板来应用统一格式（如字体、页边距）。默认True。"
                    }
                },
                "required": ["content", "output_path"]
            }
        }
    },

    # ==================== 联网搜索工具 ====================
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": (
                "按网络设置中的优先顺序使用 Tavily、知乎搜索、内建聚合或浏览器搜索获取实时信息。"
                "当用户询问新闻、天气、实时事件、最新资讯、需要查找资料或无法用本地知识回答的问题时，调用此工具。"
                "返回带来源域名和原始链接的搜索证据；摘要不是网页正文，重要结论应继续读取链接。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "要搜索的关键词，例如 '今日新闻' 或 'Python 教程'"
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "最多返回几条结果，默认 5，最大 10"
                    }
                },
                "required": ["query"]
            }
        }
    },

    # ==================== 网页内容提取工具 ====================
    {
        "type": "function",
        "function": {
            "name": "fetch_webpage",
            "description": (
                "获取指定网页的文本内容（去除HTML标签，提取主要文字）。"
                "当用户要求查看某个链接的内容、阅读某篇文章、提取网页信息时调用此工具。"
                "注意：有些网站可能限制爬取，返回内容可能不完整。"
                "如果返回 403 或内容为空，请依次尝试以下工具："
                "fetch_webpage → fetch_webpage_via_api → fetch_webpage_browser → fetch_webpage_stealth。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "要提取内容的网页URL，必须以 http:// 或 https:// 开头"
                    },
                    "max_length": {
                        "type": "integer",
                        "description": "返回的最大字符数，默认3000，避免回复过长"
                    }
                },
                "required": ["url"]
            }
        }
    },

    # ==================== 浏览器内核网页提取工具 ====================
    {
        "type": "function",
        "function": {
            "name": "fetch_webpage_browser",
            "description": (
                "使用真实浏览器内核获取网页文本内容（绕过反爬）。"
                "当普通 fetch_webpage 失败（返回 403 或内容为空）时，可以尝试调用此工具。"
                "注意：速度较慢但更可靠，适用于百度百科等严格反爬的网站。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "要提取内容的网页URL"
                    },
                    "max_length": {
                        "type": "integer",
                        "description": "返回的最大字符数，默认3000"
                    }
                },
                "required": ["url"]
            }
        }
    },

    # ==================== 通过 Jina Reader API 解析网页（绕过反爬） ====================
    {
        "type": "function",
        "function": {
            "name": "fetch_webpage_via_api",
            "description": (
                "使用 r.jina.ai 解析服务获取网页的文本内容（Markdown 格式）。"
                "当普通工具获取失败（如返回 403 或空内容）时，应优先使用此工具。"
                "适用于知乎、百度百科等反爬严格的网站。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "要解析的网页URL"
                    },
                    "max_length": {
                        "type": "integer",
                        "description": "返回的最大字符数，默认3000"
                    }
                },
                "required": ["url"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_webpage_stealth",
            "description": (
                "使用 curl_cffi TLS 指纹伪装技术获取网页文本内容（绕过反爬，速度较快）。"
                "当 fetch_webpage 和 fetch_webpage_via_api 都返回 403 或空内容时，"
                "应优先尝试此工具再尝试 fetch_webpage_browser。"
                "它能模拟 Chrome 浏览器的 TLS 握手特征，"
                "能绕过大部分基于 TLS 指纹检测的反爬机制（如 Cloudflare），"
                "比启动完整浏览器内核快得多。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "要提取内容的网页URL，必须以 http:// 或 https:// 开头"
                    },
                    "max_length": {
                        "type": "integer",
                        "description": "返回的最大字符数，默认3000，避免回复过长"
                    }
                },
                "required": ["url"]
            }
        }
    },

    # ==================== 精确文件编辑工具 ====================
    {
        "type": "function",
        "function": {
            "name": "edit_file",
            "description": (
                "在文件中精确替换指定的文字内容，不影响文件其他部分。"
                "【重要】修改文件前必须先用 read_file 读取确认内容存在，再调用此工具。"
                "适用于：修改代码中某一行、更新配置项、替换文档中某段话。"
                "比 write_file 更安全，因为只改动指定部分，不会覆盖整个文件。"
                "如果找不到 old_string，会返回错误，请检查内容是否完全一致（包括空格和换行）。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "要修改的文件的绝对路径"
                    },
                    "old_string": {
                        "type": "string",
                        "description": "要被替换掉的原始文字（必须与文件中的内容完全一致，包括空格、换行、缩进）"
                    },
                    "new_string": {
                        "type": "string",
                        "description": "用来替换的新内容"
                    },
                    "replace_all": {
                        "type": "boolean",
                        "description": "是否替换文件中所有匹配项，默认 false（只替换第一处）"
                    }
                },
                "required": ["path", "old_string", "new_string"]
            }
        }
    },

    # ==================== 文件内容搜索工具 ====================
    {
        "type": "function",
        "function": {
            "name": "grep_file",
            "description": (
                "在指定文件的内容中搜索关键词，返回所有匹配行及其行号，并显示上下文。"
                "适用于：在代码中找某个函数/变量、在文档中找某段话、确认某内容是否存在。"
                "比 read_file 更高效：大文件不必全部阅读，直接定位到关键位置。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "要搜索的文件的绝对路径"
                    },
                    "keyword": {
                        "type": "string",
                        "description": "要搜索的关键词（大小写不敏感）"
                    },
                    "context_lines": {
                        "type": "integer",
                        "description": "每个匹配行前后显示几行上下文，默认 2"
                    }
                },
                "required": ["path", "keyword"]
            }
        }
    },

    # ==================== 按行范围读取文件工具 ====================
    {
        "type": "function",
        "function": {
            "name": "read_file_lines",
            "description": (
                "读取文件中指定行范围的内容，返回带行号的文字。"
                "适用于：已知某功能在第几行附近、只需查看文件某一段、配合 grep_file 定位后精确阅读。"
                "比 read_file 更精准，不受 15000 字符分块限制。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "文件的绝对路径"
                    },
                    "start_line": {
                        "type": "integer",
                        "description": "起始行号（从 1 开始）"
                    },
                    "end_line": {
                        "type": "integer",
                        "description": "结束行号（含）。不传则读到文件末尾"
                    }
                },
                "required": ["path", "start_line"]
            }
        }
    },

    # ==================== 文件模式匹配工具 ====================
    {
        "type": "function",
        "function": {
            "name": "glob_files",
            "description": (
                "在指定目录中按文件名模式批量查找文件，支持通配符。"
                "适用于：找出所有 Python 文件、找所有 .docx 文档、找特定前缀的文件。"
                "pattern 示例：'**/*.py'（所有Python文件）、'*.txt'（当前目录txt文件）、'报告*.docx'（报告开头的Word文档）。"
                "支持精确的文件名模式匹配（通配符），比关键词搜索更精准。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "directory": {
                        "type": "string",
                        "description": "要搜索的根目录路径，默认用户主目录"
                    },
                    "pattern": {
                        "type": "string",
                        "description": "文件名匹配模式，支持 * ? ** 通配符，例如 '**/*.py' 或 '*.txt'"
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "最多返回的结果数量，默认 50"
                    }
                },
                "required": ["pattern"]
            }
        }
    },

    {
        "type": "function",
        "function": {
            "name": "ocr_image",
            "description": "仅用于提取图片中的文字（OCR）。当用户明确要求提取图片中的文字时使用。不能理解图像内容或描述画面。",
            "parameters": {
                "type": "object",
                "properties": {
                    "image_path": {
                        "type": "string",
                        "description": "图片文件的绝对路径"
                    },
                    "language": {
                        "type": "string",
                        "description": "语言，可选值：chi_sim+eng（简体中文+英文）、eng（英文），默认 chi_sim+eng"
                    }
                },
                "required": ["image_path"]
            }
        }
    },


    # ==================== 批量OCR工具 ====================
    {
        "type": "function",
        "function": {
            "name": "ocr_batch",
            "description": (
                "批量识别指定文件夹内所有图片文件（png/jpg/bmp等）中的文字，将每个文件的识别结果汇总返回。"
                "当用户要求'批量提取'、'遍历文件夹'、'把里面所有图片的文字都提取出来'时，使用此工具，不要使用 run_command。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "folder_path": {
                        "type": "string",
                        "description": "存放图片的文件夹路径"
                    },
                    "language": {
                        "type": "string",
                        "description": "语言，默认 chi_sim+eng (中英文)"
                    }
                },
                "required": ["folder_path"]
            }
        }
    },

    {
        "type": "function",
        "function": {
            "name": "describe_image",
            "description": (
                "理解并描述图片的内容——画面里有什么人、什么物体、什么场景、发生了什么动作。"
                "当用户发送了一张图片，或要求你看看某张图片的内容时，必须调用此工具来获取视觉描述。"
                "注意：已经内置的 OCR（光学字符识别）工具是专门用来提取图片中的文字的，如果用户明确只要求提取文字，优先使用 ocr_image。"
                "此工具可以看到更丰富的内容，比如：人物动作、表情、物体位置关系、场景氛围、颜色等。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "image_path": {
                        "type": "string",
                        "description": "图片文件的路径。用户发送图片时系统会自动将路径传入。"
                    },
                    "prompt": {
                        "type": "string",
                        "description": "可选的提问文本，例如 '这只猫是什么颜色？'。不需要则留空。"
                    }
                },
                "required": ["image_path"]
            }
        }
    },

    {
        "type": "function",
        "function": {
            "name": "capture_from_camera",
            "description": "从USB摄像头拍一张照片，返回保存的图片路径。",
            "parameters": {"type": "object", "properties": {}}
        }
    },

    {
        "type": "function",
        "function": {
            "name": "capture_desktop",
            "description": "截取当前电脑屏幕，返回保存的图片路径。当用户要求看看他在干什么、看看屏幕、看看桌面时调用。",
            "parameters": {"type": "object", "properties": {}}
        }
    },

    {
        "type": "function",
        "function": {
            "name": "send_file_to_qq",
            "description": (
                "将本地文件发送到主人的QQ上。当主人明确要求'把文件发到QQ'、'发给我'、"
                "'把这个发到我的QQ上'时使用。支持任何文件类型（docx/txt/图片等）。"
                "注意：只能发给主人，不能发给其他QQ好友。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "要发送的文件的绝对路径"
                    },
                    "name": {
                        "type": "string",
                        "description": "QQ上显示的文件名（可选，不传则用原文件名）"
                    }
                },
                "required": ["path"]
            }
        }
    },

    # ==================== 主动聊天开关 ====================
    {
        "type": "function",
        "function": {
            "name": "toggle_proactive_chat",
            "description": (
                "开启或关闭QQ主动聊天功能。"
                "当用户说'开启主动聊天'、'打开主动聊天'、'启用QQ主动聊天'时，设置 action='enable'。"
                "当用户说'关闭主动聊天'、'停用主动聊天'、'禁用主动聊天'时，设置 action='disable'。"
                "此操作会立即生效，无需重启。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["enable", "disable"],
                        "description": "开启(enable)或关闭(disable)主动聊天"
                    }
                },
                "required": ["action"]
            }
        }
    },

    # ==================== 统一会话历史搜索 ====================
    {
        "type": "function",
        "function": {
            "name": "search_conversation_history",
            "description": (
                "搜索真实聊天记录。当用户问最近、昨天、之前聊了什么时必须使用。"
                "mode=recent 按消息实际时间回顾，不需要关键词；mode=keyword 搜索具体内容。"
                "只有用户明确指定QQ或桌面端时才限制channels，否则搜索主人全部授权端。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "关键词；recent模式可留空"},
                    "mode": {
                        "type": "string", "enum": ["recent", "keyword"],
                        "description": "最近回顾用recent，查具体内容用keyword，默认recent"
                    },
                    "time_range": {
                        "type": "string", "enum": ["today", "yesterday", "7d", "30d", "all"],
                        "description": "时间范围；最近默认7d"
                    },
                    "channels": {
                        "type": "array",
                        "items": {"type": "string", "enum": ["desktop", "qq_private", "qq_group", "wechat_private", "wechat_group"]},
                        "description": "可选来源端；未明确指定端时不要填写"
                    },
                    "limit": {"type": "integer", "description": "最多返回消息数，默认20，最大50"}
                },
                "additionalProperties": False
            }
        }
    },

    # ==================== 近期互动联系人聚合（仅主人） ====================
    {
        "type": "function",
        "function": {
            "name": "query_recent_contacts",
            "description": (
                "查询最近与莲心聊过天的其他用户（非主人），供主人回顾。"
                "当主人问『最近有什么人找过你/跟你聊过天/都聊了什么』时必须使用。"
                "返回每位联系人的身份标识、渠道、最后活跃时间和最近几条消息。"
                "此工具仅对主人会话开放，其它用户无法调用。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "days": {"type": "integer", "description": "回顾最近N天，默认7，最大365"},
                    "per_contact_limit": {"type": "integer", "description": "每位联系人返回最近消息条数，默认3，最大10"},
                    "max_contacts": {"type": "integer", "description": "最多返回几位联系人，默认10，最大50"}
                },
                "additionalProperties": False
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "query_qq_friend_list",
            "description": (
                "查询莲心绑定 QQ 账号的好友列表（昵称、QQ号、备注），供主人回顾。"
                "当主人问『你有哪些QQ好友/你有几个好友/我的QQ好友都有谁』时必须使用。"
                "默认使用最近缓存；传 refresh=true 可强制向 QQ 重新拉取最新列表。"
                "此工具仅对主人会话开放，其它用户无法调用。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "refresh": {"type": "boolean", "description": "是否强制刷新最新好友列表，默认 false 使用缓存"},
                    "keyword": {"type": "string", "description": "按昵称/备注/QQ号关键词筛选，可省略"}
                },
                "additionalProperties": False
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "configure_network_tools",
            "description": (
                "查看或调整联网搜索/网页读取工具的启停状态和优先顺序。"
                "修改配置属于副作用操作，只有用户明确要求时才能调用；普通联网请求不要调用。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["status", "enable", "disable", "move", "reset"],
                        "description": "查看、启用、禁用、移动或恢复默认"
                    },
                    "kind": {
                        "type": "string", "enum": ["search", "fetch"],
                        "description": "搜索工具或网页读取工具"
                    },
                    "tool_id": {"type": "string", "description": "工具 ID；status/reset 可省略"},
                    "position": {"type": "integer", "description": "move 的目标位置，从 0 开始"}
                },
                "required": ["action"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "clear_document_cache",
            "description": "清理莲心为 DOCX/PDF 生成的 Markdown 缓存。仅当用户明确要求清理文档缓存时使用；不会删除或修改任何原始文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "confirm": {
                        "type": "boolean",
                        "description": "仅在用户已明确确认清理时传 true"
                    }
                },
                "required": ["confirm"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "update_current_state",
            "description": (
                "维护用户有时效的当前状态，而不是永久记忆。适用于生病、情绪低落、"
                "出差地点、正在推进的项目、短期计划等会变化或会过期的信息。"
                "action=set 新建状态；update 更新已知状态；resolve 明确结束状态；"
                "list 查看当前有效状态。用户只是在表达稳定偏好或长期事实时不要调用。"
                "不要根据含糊表达推断医学诊断；不确定的信息必须标记为 inferred。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["set", "update", "resolve", "list"],
                        "description": "要执行的状态操作"
                    },
                    "state_id": {
                        "type": "integer",
                        "description": "update/resolve 时使用的状态编号"
                    },
                    "content": {
                        "type": "string",
                        "description": "具体且可独立理解的状态描述"
                    },
                    "state_type": {
                        "type": "string",
                        "enum": ["health", "emotion", "location", "project", "relationship", "plan", "other"],
                        "description": "状态类型"
                    },
                    "expires_at": {
                        "type": "string",
                        "description": "可选的 ISO 8601 绝对过期时间"
                    },
                    "duration_days": {
                        "type": "integer",
                        "minimum": 1,
                        "maximum": 90,
                        "description": "未给绝对时间时，从现在起有效的天数"
                    },
                    "confidence": {
                        "type": "number",
                        "minimum": 0,
                        "maximum": 1,
                        "description": "对状态判断的置信度"
                    },
                    "source_quality": {
                        "type": "string",
                        "enum": ["direct_statement", "user_confirmed", "inferred"],
                        "description": "信息来自用户直述、用户确认或模型推断"
                    },
                    "resolve_reason": {
                        "type": "string",
                        "description": "resolve 时说明状态为何结束"
                    }
                },
                "required": ["action"],
                "additionalProperties": False
            }
        }
    },

    # ==================== 旧跨端搜索兼容工具 ====================
    {
        "type": "function",
        "function": {
            "name": "search_cross_session",
            "description": (
                "搜索另一端（桌面端↔QQ端）的历史聊天记录。"
                "当用户询问'之前在电脑上聊了什么'、'回忆一下QQ上说过的话'、"
                "'帮我找找另一边之前提到的内容'等涉及另一端聊天回忆的问题时，调用此工具。"
                "注意：搜索的是另一端的全部历史记录，不限最近几条。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "keyword": {
                        "type": "string",
                        "description": "要搜索的关键词，例如'火锅'、'项目'、'Python'等"
                    },
                    "limit": {
                        "type": "integer",
                        "description": "最多返回几条匹配结果，默认 5，最大 10"
                    }
                },
                "required": ["keyword"]
            }
        }
    },

    # ==================== 技能系统工具 ====================
    {
        "type": "function",
        "function": {
            "name": "list_skills",
            "description": (
                "仅当用户明确询问'你有什么技能''有哪些技能''激活了什么技能'时才调用此工具。"
                "不要在每次对话开始时调用。不要在文件操作、搜索、读取、对比等非技能相关任务中调用。"
                "\u26a0\ufe0f 文件操作（读取、搜索、对比、编辑）不需要技能，直接使用 read_file/search_files_everything/diff_files 等工具即可。"
            ),
            "parameters": {
                "type": "object",
                "properties": {},
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "activate_skill",
            "description": (
                "激活指定的技能包。技能包激活后可获得额外能力或知识。"
                "当用户要求激活某个技能时（如'启动XX技能'），先调用 list_skills 查看可用技能，再调用此工具激活。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {
                        "type": "string",
                        "description": "要激活的技能名称，例如 '文本工具'"
                    }
                },
                "required": ["name"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "deactivate_skill",
            "description": "停用指定的技能包。当用户要求关闭某个技能时调用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {
                        "type": "string",
                        "description": "要停用的技能名称"
                    }
                },
                "required": ["name"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "search_memory",
            "description": (
                "在长期记忆中搜索包含指定关键词的事实。"
                "当用户问'你还记得关于XX的事吗'、'我之前是不是跟你说过XX'、"
                "'翻一下我的记忆，关于XX的'等涉及回忆具体内容的问题时，调用此工具。"
                "注意：长期记忆存储的是你（莲心）保存下来的事实，"
                "不是聊天记录。如果要搜聊天记录，请使用 search_cross_session。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "keyword": {
                        "type": "string",
                        "description": "要搜索的关键词，例如'毕业'、'游戏'、'生日'等"
                    },
                    "category": {
                        "type": "string",
                        "enum": ["profile", "preferences", "events", "knowledge", "behaviors", "skills"],
                        "description": "可选：限定搜索的分类，不填则搜索所有分类"
                    }
                },
                "required": ["keyword"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "trace_memory_source",
            "description": (
                "追溯一条长期记忆的证据来源。当用户问‘你为什么记得’、"
                "‘这条记忆来自哪次对话’或质疑某条记忆时调用。"
                "memory_id 来自 search_memory/RAG 结果中的‘记忆#编号’。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "memory_id": {
                        "type": "integer",
                        "description": "要追溯的长期记忆编号"
                    }
                },
                "required": ["memory_id"],
                "additionalProperties": False
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "explain_memory_quality",
            "description": (
                "解释一条长期记忆的质量评分、证据数量、来源可靠性、新鲜度、召回次数"
                "和当前复核状态。用户质疑莲心为什么保留或召回某条记忆时调用。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "memory_id": {
                        "type": "integer", "description": "记忆编号"
                    }
                },
                "required": ["memory_id"],
                "additionalProperties": False
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "review_memory_conflict",
            "description": (
                "对长期记忆冲突候选进行语义裁决。代码只提供候选，相似度不能代替语义判断。"
                "action=list 查看待裁决候选；action=resolve 时判断新记忆相对于旧记忆是："
                "duplicate（重复）、complements（补充）、contradicts（矛盾但不能确定谁取代谁）、"
                "supersedes（新事实明确取代旧事实）或 unrelated（无关）。"
                "只有语义和时间关系明确时才能选择 supersedes。"
                "工具返回会以【合并已执行】/【替换已执行】/【未合并·未执行】/【裁决失败·未执行】/【裁决已记录·未删除】标明真实状态，请严格按返回状态向用户汇报；未执行或失败时不得声称已合并。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string", "enum": ["list", "resolve"],
                        "description": "查看候选或提交语义裁决"
                    },
                    "candidate_id": {
                        "type": "integer", "description": "resolve 时使用的冲突候选编号"
                    },
                    "decision": {
                        "type": "string",
                        "enum": ["duplicate", "complements", "contradicts", "supersedes", "unrelated"],
                        "description": "新记忆相对于旧记忆的语义关系"
                    },
                    "confidence": {
                        "type": "number", "minimum": 0, "maximum": 1,
                        "description": "本次语义裁决的置信度"
                    },
                    "rationale": {
                        "type": "string",
                        "description": "基于两条事实语义和时间关系的简短理由"
                    }
                },
                "required": ["action"],
                "additionalProperties": False
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "update_memory",
            "description": (
                "更新或覆盖长期记忆中已存在的事实。"
                "当用户说'之前说的那个不对，其实是这样的'、'改一下，不是XX，是XX'、"
                "'我要纠正一下之前的说法'时，使用此工具更新对应的事实。"
                "会搜索所有包含 old_keyword 的事实，将它们替换为 new_fact。"
                "如果找不到匹配的旧事实，则直接作为新事实保存。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "old_keyword": {
                        "type": "string",
                        "description": "要更新的旧事实中的关键词，用于定位需要更新的事实。例如旧事实是'雨心的生日是4月25日'，想把它替换时，old_keyword 可以填'生日'"
                    },
                    "new_fact": {
                        "type": "string",
                        "description": "更新后的事实内容，例如'雨心的生日是4月26日'。要求信息完整、一句话说清楚"
                    },
                    "category": {
                        "type": "string",
                        "enum": ["profile", "preferences", "events", "knowledge", "behaviors", "skills"],
                        "description": "可选：限定搜索的分类，不填则搜索所有分类"
                    }
                },
                "required": ["old_keyword", "new_fact"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "delete_memory",
            "description": (
                "从长期记忆中删除指定的事实。"
                "当用户说'忘掉关于XX的事'、'删掉那条记忆'、'那个信息现在没用了删掉吧'时，"
                "搜索包含 keyword 的事实并删除所有匹配项。"
                "注意：删除是不可恢复的操作，删除前应在回复中告知用户将被删除的内容。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "keyword": {
                        "type": "string",
                        "description": "要删除的事实中的关键词。所有包含此关键词的事实都会被删除"
                    }
                },
                "required": ["keyword"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "set_expression",
            "description": (
                "切换莲心的立绘表情。当你想通过表情来表达情绪时使用，"
                "比如开心、生气、伤心、惊讶等。"
                "此工具仅负责切换显示图片，不会影响对话内容。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "emotion": {
                        "type": "string",
                        "description": "情绪/表情名称，如 开心、生气、伤心、惊讶、疑惑、害羞、撒娇、疲惫、默认"
                    }
                },
                "required": ["emotion"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "list_memories",
            "description": "查看长期记忆中保存的全部内容，按分类展示。当用户说\"你都记住了什么\"、\"翻一下我的记忆\"、\"让我看看你记了哪些东西\"时调用。",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "search_graph_memory",
"description": "统一搜索长期记忆：同时搜索分类事实和知识图谱实体关系。比 search_memory 更全面——一次调用返回分类事实 + 实体关系边。适用于所有需要回顾信息的问题，尤其是涉及「谁」「什么关系」「哪些关联」的查询。必填参数名是 keywords（字符串数组），从用户问题提取实体名和关系词填入 keywords；不要使用 query 或 keyword 作为参数名。",

            "parameters": {
                "type": "object",
                "properties": {
                    "keywords": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "搜索关键词列表。从用户问题中提取所有相关实体名和关系词。例如问'我跟小明喜欢什么音乐'应提取['我','小明','喜欢','音乐']；问'上次聊的项目用了哪些技术'应提取['项目','技术']"
                    },
                    "entity_type": {
                        "type": "string",
                        "enum": ["人物", "地点", "组织", "物品", "概念", "时间", "事件", "活动", "技术", "文件"],
                        "description": "可选，限定实体类型"
                    }
                },
                "required": ["keywords"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "query_connected_entities",
            "description": "查找与某个实体间接关联的所有实体和关系（图谱多跳遍历）。适用于用户问\"这个项目用了哪些技术\"、\"我的朋友都住在哪里\"、\"和X相关的所有信息\"。",
            "parameters": {
                "type": "object",
                "properties": {
                    "entity_name": {
                        "type": "string",
                        "description": "要查询的实体名称"
                    },
                    "depth": {
                        "type": "integer",
                        "description": "查询深度（1=直接关联, 2=两跳, 最多3），默认1"
                    }
                },
                "required": ["entity_name"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "delete_graph_entity",
            "description": "从知识图谱中删除指定实体及其所有关联边（不可恢复）。适用于删除错误或测试数据。",
            "parameters": {
                "type": "object",
                "properties": {
                    "entity_name": {
                        "type": "string",
                        "description": "要删除的实体名称（精确匹配）"
                    }
                },
                "required": ["entity_name"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "add_graph_edge",
            "description": (
                "手动向知识图谱添加一条五元组边（实体→关系→实体）。"
                "适合将已知的事实关系写入图记忆，例如从分类记忆中提取的关系。"
                "如果边已存在，会增强其强度而非重复添加。"
                "实体类型必须是：人物/地点/组织/物品/概念/时间/事件/活动/技术/文件"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "head": {"type": "string", "description": "主体实体名，如：雨心"},
                    "head_type": {"type": "string", "description": "主体实体类型"},
                    "relation": {"type": "string", "description": "关系名，如：工作于、喜欢、创造"},
                    "tail": {"type": "string", "description": "客体实体名，如：润建公司"},
                    "tail_type": {"type": "string", "description": "客体实体类型"}
                },
                "required": ["head", "head_type", "relation", "tail", "tail_type"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "remove_graph_edge",
            "description": (
                "从知识图谱中删除一条指定的边（不影响实体本身）。"
                "适用于纠正错误关系或清理过时信息。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "head": {"type": "string", "description": "主体实体名"},
                    "relation": {"type": "string", "description": "关系名"},
                    "tail": {"type": "string", "description": "客体实体名"}
                },
                "required": ["head", "relation", "tail"]
            }
        }
    },

    
    # ==================== 天气工具 ====================
    {
        "type": "function",
        "function": {
            "name": "get_weather",
            "description": (
                "查询指定城市的实时天气和天气预报。支持实时(current)、逐小时(hourly)、"
                "逐天(daily)、完整(full)四种模式。当用户问天气、气温、多少度、冷不冷、"
                "热不热、下雨、刮风、下雪等天气相关问题时调用。"
                "如果用户没有说城市名，莲心会先使用配置里设置的默认城市（如广州）；"
                "若未设置默认城市，再从记忆中查找，仍找不到再反问用户。"
                "获取天气后会附带出行建议（带伞/加衣/防晒等）。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "city": {
                        "type": "string",
                        "description": "城市名称，如'北京'、'上海'、'广州'。如果用户没有说城市名，此参数可留空，莲心会使用配置里的默认城市（如广州），或从记忆里读取。"
                    },
                    "forecast_type": {
                        "type": "string",
                        "enum": ["current", "hourly", "daily", "full"],
                        "description": "查询模式：current(实时天气)、hourly(逐小时预报)、daily(逐天预报)、full(实时+3天预报+出行建议，默认)"
                    }
                },
                "required": ["city"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "set_user_city",
            "description": (
                "设置或更新用户的所在城市。当用户告诉你'我在XX'、'我住在XX'、"
                "'我的城市是XX'、'我在XX上学/工作'、'我目前在XX'时调用此工具。"
                "保存后莲心就能记住用户所在的城市，以后查询天气时自动使用这个城市，"
                "无需用户每次都说城市名。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "city": {
                        "type": "string",
                        "description": "用户所在城市名称，如'北京'、'上海'、'广州'"
                    }
                },
                "required": ["city"]
            }
        }
    },
        # ── 第一阶段新增：编程工具增强 ─────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "search_code",
            "description": (
                "在目录中搜索代码内容（正则表达式匹配），支持多文件、上下文行。\n"
                "比 grep_file 更强大：支持文件类型过滤、大小写不敏感、上下文行显示。\n"
                "用于：查找函数定义、变量引用、TODO 注释、跨文件模式搜索。\n"
                "提示：简单单文件搜索直接用 grep_file 更快。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "pattern": {
                        "type": "string",
                        "description": "搜索的正则表达式模式，如 'def main' 或 'TODO'"
                    },
                    "directory": {
                        "type": "string",
                        "description": "搜索目录的绝对路径，默认当前工作目录"
                    },
                    "file_pattern": {
                        "type": "string",
                        "description": "文件名匹配模式（glob），如 '*.py'、'*.{js,ts}'。不指定则搜索所有文本文件"
                    },
                    "context_lines": {
                        "type": "integer",
                        "description": "显示匹配行前后的上下文行数，默认 0"
                    },
                    "case_sensitive": {
                        "type": "boolean",
                        "description": "是否大小写敏感，默认 true"
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "最多返回几条结果，默认 30"
                    },
                    "exclude_pattern": {
                        "type": "string",
                        "description": "额外排除的目录/文件 glob 模式。默认自动排除 .git、node_modules 等"
                    }
                },
                "required": ["pattern"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "diff_files",
            "description": (
                "对比两个文本文件的差异（类似 Git diff 格式）。\n"
                "用于：查看修改前后的变化、验证编辑结果、对比代码版本。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "file_a": {
                        "type": "string",
                        "description": "第一个文件的绝对路径"
                    },
                    "file_b": {
                        "type": "string",
                        "description": "第二个文件的绝对路径"
                    },
                    "context_lines": {
                        "type": "integer",
                        "description": "差异周围显示几行上下文，默认 3"
                    }
                },
                "required": ["file_a", "file_b"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "run_shell",
            "description": (
                "执行 Shell 命令并返回输出。比 run_command 更强大。\n"
                "支持：指定工作目录、超时控制、输出行数限制。\n"
                "用于：编译代码、运行测试、安装依赖、Git 操作等。\n"
                "安全注意：不会自动执行危险命令。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "command": {
                        "type": "string",
                        "description": "要执行的 Shell 命令"
                    },
                    "working_dir": {
                        "type": "string",
                        "description": "工作目录的绝对路径，默认当前目录"
                    },
                    "timeout": {
                        "type": "integer",
                        "description": "超时秒数，默认 60"
                    },
                    "max_output_lines": {
                        "type": "integer",
                        "description": "最大输出行数，超出截断。默认 200"
                    }
                },
                "required": ["command"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "git_status",
            "description": (
                "查看 Git 仓库的状态信息。\n"
                "用于：查看修改了哪些文件、当前分支、最近提交记录等。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "repo_path": {
                        "type": "string",
                        "description": "Git 仓库的绝对路径，默认当前工作目录"
                    },
                    "action": {
                        "type": "string",
                        "enum": ["status", "diff", "log", "branch"],
                        "description": "操作类型：status（文件状态）、diff（查看改动）、log（最近提交）、branch（分支列表）"
                    },
                    "limit": {
                        "type": "integer",
                        "description": "对于 log 操作，限制显示的提交数，默认 10"
                    }
                },
                "required": ["action"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "code_structure",
            "description": (
                "快速列出代码文件中定义的函数、类、方法等结构。\n"
                "支持：Python、JavaScript、TypeScript、Java、Go、Rust 等语言。\n"
                "用于：快速了解文件结构、定位函数定义位置。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "要分析的代码文件的绝对路径"
                    }
                },
                "required": ["file_path"]
            }
        }
    },
    # ── 第二阶段：子代理任务分解系统 ─────────────────────
    {
        "type": "function",
        "function": {
            "name": "plan_tasks",
            "description": (
                "将复杂任务分解为可并行执行的子任务列表。\n"
                "使用内置 LLM 分析任务，生成结构化的子任务计划。\n"
                "返回的子任务可逐个交给 delegate_task 并行执行。\n"
                "用于：复杂重构、多文件修改、多步骤操作等需要分步执行的任务。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "task_description": {
                        "type": "string",
                        "description": "要分解的复杂任务描述"
                    },
                    "context": {
                        "type": "string",
                        "description": "额外上下文信息，如项目结构、相关文件列表等"
                    },
                    "max_subtasks": {
                        "type": "integer",
                        "description": "最多拆分为几个子任务，默认 5"
                    }
                },
                "required": ["task_description"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "delegate_task",
            "description": (
                "将子任务委派给一个独立的子代理执行。子代理拥有受限的工具集，\n"
                "专注于完成单一任务，完成后返回结果。\n"
                "多个 delegate_task 调用会自动并行执行。\n\n"
                "【子代理可用工具】read_file, read_file_lines, grep_file, search_code,\n"
                "glob_files, list_directory, edit_file, code_structure, run_shell,\n"
                "run_python_code, diff_files, git_status\n\n"
                "【使用场景】\n"
                "- 同时修改多个文件时，每个文件委派一个子代理\n"
                "- 搜索分析类任务与修改类任务并行执行\n"
                "- 需要独立上下文执行的分析任务"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "task": {
                        "type": "string",
                        "description": "子代理要执行的具体任务描述，越具体越好"
                    },
                    "working_dir": {
                        "type": "string",
                        "description": "子代理的工作目录，默认继承主代理的目录"
                    },
                    "timeout_seconds": {
                        "type": "integer",
                        "description": "超时秒数，默认 120"
                    },
                    "max_iterations": {
                        "type": "integer",
                        "description": "子代理最多调用几轮工具，默认 10"
                    }
                },
                "required": ["task"]
            }
        }
    },
    # 第三阶段：任务进度追踪
    {
        "type": "function",
        "function": {
            "name": "track_tasks",
            "description": "更新当前会话的任务清单。用于追踪复杂多步骤任务的进度。"
                           "每次调用会全量替换整个列表，不是增量追加。\n"
                           "规则：同一时刻最多一个任务 in_progress，完成后立即标记 completed，"
                           "全部完成时传空列表清除。",
            "parameters": {
                "type": "object",
                "properties": {
                    "todos": {
                        "type": "array",
                        "description": "任务列表，全量替换。空列表表示清除所有任务。",
                        "items": {
                            "type": "object",
                            "properties": {
                                "content": {
                                    "type": "string",
                                    "description": "任务描述（祈使句，如'添加用户登录功能'）"
                                },
                                "status": {
                                    "type": "string",
                                    "enum": ["pending", "in_progress", "completed"],
                                    "description": "任务状态：pending=待做 in_progress=进行中 completed=已完成"
                                },
                                "activeForm": {
                                    "type": "string",
                                    "description": "进行时描述（如'正在添加用户登录功能...'）"
                                }
                            },
                            "required": ["content", "status", "activeForm"]
                        }
                    }
                },
                "required": ["todos"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "code_goto_def",
            "description": "跳转到 Python 函数/类/变量的定义位置。给定文件和行号，精确返回定义的文件路径、行号和代码。支持跨文件追踪（如导入的函数）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "源文件路径，如 'brain/agent.py'"},
                    "line": {"type": "integer", "description": "光标所在行号"},
                    "symbol": {"type": "string", "description": "要查找的符号名（可选，不传则根据行号自动推断）"},
                    "column": {"type": "integer", "description": "光标所在列号（可选，默认 0）"}
                },
                "required": ["file_path", "line"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "code_find_refs",
            "description": "查找 Python 函数/类/变量在项目中的所有引用位置（调用处、赋值处等），返回精确的文件路径和行号列表。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "源文件路径"},
                    "line": {"type": "integer", "description": "符号定义所在行号"},
                    "symbol": {"type": "string", "description": "要查找引用的符号名（可选）"},
                    "column": {"type": "integer", "description": "符号定义所在列号（可选，默认 0）"}
                },
                "required": ["file_path", "line"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "code_diagnostics",
            "description": "检查 Python 文件的语法错误、未使用变量、重复定义等问题。修改代码后建议调用此工具验证。",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "要检查的 Python 文件路径"}
                },
                "required": ["file_path"]
            }
        }
    },

    # ==================== 图片生成工具 ====================
    {
        "type": "function",
        "function": {
            "name": "generate_image",
            "description": (
                "根据文字描述生成图片。使用 Agnes Image API，支持多种尺寸和质量。"
                "当用户要求画图、生成图片、创作图像、制作插图时调用此工具。"
                "生成后图片会自动保存，用户可以查看和下载。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": "图片描述（中文或英文），越详细越好。例如：'一只可爱的橘猫坐在窗台上，窗外是星空，油画风格'"
                    },
                    "size": {
                        "type": "string",
                        "enum": ["1024x1024", "1792x1024", "1024x1792", "4k"],
                        "description": "图片尺寸，默认使用配置中的默认值"
                    },
                    "quality": {
                        "type": "string",
                        "enum": ["standard", "hd"],
                        "description": "图片质量，默认使用配置中的默认值"
                    }
                },
                "required": ["prompt"]
            }
        }
    },

    # ==================== 视频生成工具 ====================
    {
        "type": "function",
        "function": {
            "name": "generate_video",
            "description": (
                "根据文字或图片生成视频。使用 Agnes Video API，支持文生视频和图生视频。"
                "视频生成是异步任务，约需 1-5 分钟完成。"
                "当用户要求生成视频、制作动画、让图片动起来时调用此工具。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": "视频内容描述（中文或英文）。图生视频时描述需要哪些内容运动。"
                    },
                    "image_url": {
                        "type": "string",
                        "description": "参考图片的 URL（可选）。如果提供，将基于该图片生成视频（图生视频模式）。"
                    },
                    "duration": {
                        "type": "integer",
                        "description": "视频时长（秒），可选 3、5、10、18。默认 5 秒。"
                    }
                },
                "required": ["prompt"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "bilibili_search",
            "description": (
                "在B站（bilibili.com）搜索视频。根据关键词搜索相关视频，"
                "返回视频标题、作者、播放量、BV号和链接。"
                "当用户要求搜索B站视频、找B站内容、推荐视频时使用。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "keyword": {
                        "type": "string",
                        "description": "搜索关键词，如「口琴教程」「赛博朋克」「猫猫」"
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "最多返回的结果数，默认10"
                    }
                },
                "required": ["keyword"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "bilibili_add_tag",
            "description": (
                "添加一个B站搜索兴趣标签。当用户说「以后帮我搜XXX」「我想看XXX的视频」"
                "「帮我关注XXX」时，调用此工具将关键词添加到莲心的兴趣标签库。"
                "莲心会在空闲时自动用这些标签去B站搜索视频推荐给用户。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "keyword": {
                        "type": "string",
                        "description": "兴趣标签关键词，如「口琴」「赛博朋克」「猫猫」"
                    }
                },
                "required": ["keyword"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "bilibili_list_tags",
            "description": (
                "查看当前所有B站搜索兴趣标签。当用户问「你关注哪些兴趣」「"
                "你有哪些B站标签」「帮我看看标签」时使用。"
            ),
            "parameters": {
                "type": "object",
                "properties": {},
                "required": []
            }
        }
    },


]

# ── 工具执行函数 ─────────────────────────────────────────────

def read_file(path: str) -> str:
    """读取文件第 0 块内容（开头 _CHUNK_SIZE 字符）。"""
    if str(path or "").lower().startswith(("http://", "https://")):
        # 防御性兼容：正常情况下请求级白名单不会让模型对 URL 调用
        # read_file；旧会话或第三方模型仍这样调用时，直接走统一网络路由。
        return fetch_webpage(path)
    return read_file_chunk(path, chunk_index=0)


def read_file_chunk(path: str, chunk_index: int = 0) -> str:
    """读取文件的指定分块，每块最多 _CHUNK_SIZE 字符。"""
    try:
        p = Path(path)
        if not p.exists():
            return f"错误：文件不存在 → {path}"
        if not p.is_file():
            # 目录路径 → 自动降级为列出目录内容，而不是报错让模型重试
            return f"⚠️ 「{path}」是一个目录，不是文件。已自动列出目录内容：\n{list_directory(path)}"

        content, err = _extract_full_text(p)
        if err:
            return f"读取文件出错: {err}"
        if not content:
            return "（文件内容为空）"

        total_chars  = len(content)
        total_chunks = max(1, (total_chars + _CHUNK_SIZE - 1) // _CHUNK_SIZE)

        if chunk_index < 0 or chunk_index >= total_chunks:
            return (
                f"错误：chunk_index={chunk_index} 超出范围。"
                f"该文件共 {total_chunks} 块（0 ~ {total_chunks-1}）。"
            )

        start = chunk_index * _CHUNK_SIZE
        end   = start + _CHUNK_SIZE
        chunk = content[start:end]

        header = (
            f"[文件：{p.name} | 第 {chunk_index+1}/{total_chunks} 块 "
            f"| 字符 {start+1}~{min(end, total_chars)}/{total_chars}]\n"
            + "─" * 50 + "\n"
        )
        footer = ""
        if chunk_index < total_chunks - 1:
            footer = (
                f"\n\n… [本块结束。文件共 {total_chunks} 块，"
                f"如需继续阅读请调用 read_file_chunk(path, chunk_index={chunk_index+1})]"
            )

        return header + chunk + footer

    except Exception as e:
        return f"读取文件出错: {e}"


def _extract_full_text(p: Path) -> tuple[str, str]:
    """
    提取文件完整文本，不做长度截断。
    返回 (content, error_message)，成功时 error_message 为空字符串。
    """
    ext = p.suffix.lower()
    try:
        if ext == ".docx":
            return _extract_docx_markdown(p), ""
        elif ext == ".doc":
            return _extract_doc(p), ""
        elif ext == ".pdf":
            return _extract_pdf_markdown(p), ""
        elif ext in (".xlsx", ".xls"):
            return _extract_xlsx(p), ""
        else:
            return _extract_text(p), ""
    except Exception as e:
        return "", str(e)


def _convert_with_markitdown(p: Path) -> str:
    """Use MarkItDown to preserve document structure as Markdown."""
    try:
        from markitdown import MarkItDown
    except ImportError as exc:
        raise RuntimeError(
            "未安装 MarkItDown，请执行：pip install 'markitdown[docx,pdf]==0.1.6'"
        ) from exc

    result = MarkItDown().convert(str(p))
    content = (getattr(result, "text_content", "") or "").strip()
    if not content:
        raise RuntimeError("MarkItDown 未提取到有效文本")
    return content


def _get_document_markdown_cache() -> MarkdownDocumentCache:
    """Create the process-local cache lazily so startup has no cache I/O."""
    global _document_cache
    if _document_cache is None:
        with _document_cache_lock:
            if _document_cache is None:
                _document_cache = MarkdownDocumentCache()
    return _document_cache


def _convert_with_markitdown_cached(p: Path) -> str:
    """Convert once per file content and reuse private Markdown cache entries."""
    return _get_document_markdown_cache().get_or_create(
        p, _convert_with_markitdown
    ).content


def clear_document_cache(confirm: bool = False) -> str:
    """Clear only generated Markdown cache entries after explicit confirmation."""
    if not confirm:
        return "请在确认要清理莲心的文档 Markdown 缓存后再执行。原始 DOCX/PDF 不会被删除。"
    cache = _get_document_markdown_cache()
    entries, total_bytes = cache.stats()
    cache.clear()
    return f"已清理 {entries} 份文档 Markdown 缓存（{total_bytes} 字节）。原始文件未被修改。"


def _extract_docx_markdown(p: Path) -> str:
    """Convert DOCX to Markdown, falling back to the legacy extractor."""
    try:
        return _convert_with_markitdown_cached(p)
    except Exception as markitdown_error:
        logger.warning("MarkItDown DOCX conversion failed for %s: %s", p, markitdown_error)
        try:
            content = _extract_docx(p)
        except Exception as legacy_error:
            raise RuntimeError(
                f"DOCX 转换失败（MarkItDown: {markitdown_error}; 旧转换器: {legacy_error}）"
            ) from markitdown_error
        if content.strip():
            return content
        raise RuntimeError(f"DOCX 转换失败：{markitdown_error}") from markitdown_error


def _extract_pdf_markdown(p: Path) -> str:
    """Convert PDF to Markdown, falling back to the legacy extractor."""
    try:
        return _convert_with_markitdown_cached(p)
    except Exception as markitdown_error:
        logger.warning("MarkItDown PDF conversion failed for %s: %s", p, markitdown_error)
        try:
            content = _extract_pdf(p)
        except Exception as legacy_error:
            raise RuntimeError(
                f"PDF 转换失败（MarkItDown: {markitdown_error}; 旧转换器: {legacy_error}）"
            ) from markitdown_error
        if content.strip():
            return content
        raise RuntimeError(f"PDF 转换失败：{markitdown_error}") from markitdown_error


def _extract_text(p: Path) -> str:
    """读取文本文件全文，自动检测编码（chardet 优先，回退多编码尝试）。"""
    raw = p.read_bytes()
    if not raw:
        return ""

    encoding = "utf-8"
    try:
        import chardet
        detected  = chardet.detect(raw)
        enc        = detected.get("encoding") or "utf-8"
        confidence = detected.get("confidence") or 0
        if confidence >= 0.6:
            encoding = enc
    except ImportError:
        for enc in ("utf-8", "gbk", "gb2312", "gb18030", "latin-1"):
            try:
                raw.decode(enc)
                encoding = enc
                break
            except (UnicodeDecodeError, LookupError):
                continue

    return raw.decode(encoding, errors="replace")


def _extract_docx(p: Path) -> str:
    """提取 Word .docx 文件全文（正文段落 + 表格），无长度限制。"""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("未安装 python-docx，请执行：pip install python-docx")

    doc   = Document(str(p))
    parts = []

    # 按文档中实际顺序遍历：段落和表格交替出现在 doc.element.body 里
    from docx.oxml.ns import qn
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            # 段落
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc)
            text = para.text.strip()
            if text:
                parts.append(text)
        elif tag == "tbl":
            # 表格
            from docx.table import Table
            table = Table(child, doc)
            parts.append("")   # 空行分隔
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
            parts.append("")

    return "\n".join(parts)


def _extract_doc(p: Path) -> str:
    """提取旧版 Word .doc 文件文本（Word 97-2003 格式）。

    Windows 环境下优先用 Word COM 自动化（最准确），
    次选 docx2txt / olefile，最后回退到原始文本读取。
    """
    # 1) Windows COM：调用本机 Word 打开并提取全文（准确率最高）
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        try:
            doc = word.Documents.Open(str(p.absolute()))
            text = doc.Content.Text
            doc.Close()
            if text and text.strip():
                return text.strip()
        except Exception:
            pass
        finally:
            try:
                word.Quit()
            except Exception:
                pass
    except ImportError:
        pass
    except Exception:
        pass

    # 2) 尝试 docx2txt（部分支持 .doc）
    try:
        import docx2txt
        text = docx2txt.process(str(p))
        if text and text.strip():
            return text.strip()
    except ImportError:
        pass
    except Exception:
        pass

    # 3) 尝试 olefile 提取 Word 文本流
    try:
        import olefile
        ole = olefile.OleFileIO(str(p))
        if ole.exists("WordDocument"):
            stream = ole.openstream("WordDocument")
            raw = stream.read()
            text = raw.decode("utf-16-le", errors="ignore")
            readable = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
            if readable.strip():
                return readable.strip()
    except ImportError:
        pass
    except Exception:
        pass

    # 4) 最后回退：当作纯文本尝试读取
    try:
        return _extract_text(p)
    except Exception:
        pass

    raise RuntimeError(
        "无法提取 .doc 文件内容。请确保本机安装了 Microsoft Word，\n"
        "或使用较新的 .docx 格式。"
    )


def _extract_xlsx(p: Path) -> str:
    """提取 Excel .xlsx 文件内容（所有工作表），以表格形式返回。"""
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("未安装 openpyxl，请执行：pip install openpyxl")

    wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = "\t".join(cells)
            if line.strip():
                rows.append(line)
        if rows:
            parts.append(f"[工作表: {sheet_name}]\n" + "\n".join(rows))
    wb.close()

    if not parts:
        return "（工作簿为空）"
    return "\n\n".join(parts)


def _extract_pdf(p: Path) -> str:
    """提取 PDF 全文，逐页拼接，无长度限制。"""
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("未安装 pdfplumber，请执行：pip install pdfplumber")

    parts = []
    with pdfplumber.open(str(p)) as pdf:
        total_pages = len(pdf.pages)
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                parts.append(f"[第 {i}/{total_pages} 页]\n{text.strip()}")

    if not parts:
        return ""
    return "\n\n".join(parts)


def save_memory(fact: str, category: str | None = None) -> str:
    """将事实写入长期记忆（分类存储），使用 SQLite 后端。"""
    from datetime import datetime
    from config import get_memory_config

    _ensure_migrated()

    # 未指定分类时从配置读取默认分类
    if not category:
        try:
            category = get_memory_config().get("default_save_category", "knowledge")
        except Exception:
            category = "knowledge"

    fact = fact.strip()
    if not fact:
        return "记忆内容不能为空。"
    from brain.persona.authority import is_assistant_identity_fact
    if is_assistant_identity_fact(fact):
        return (
            "这条内容描述的是莲心自身的人格或外貌设定，不能写入用户长期记忆。"
            "请在人格中枢修改当前人格档案；当前人格档案会作为唯一权威来源。"
        )

    # 自动追加记录日期
    today = datetime.now().strftime("%Y-%m-%d")
    date_tag = f"【记录于{today}】"
    if date_tag not in fact:
        fact = f"{fact} {date_tag}"

    provenance = _current_memory_provenance()
    occurred_at = provenance["occurred_at"] or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    fact_id = _memory_add(
        fact, category, source="user_saved",
        source_session_id=provenance["source_session_id"],
        source_channel=provenance["source_channel"], occurred_at=occurred_at,
    )
    if fact_id:
        _memory_add_fragment(
            fact_id, fact, category, source="user_saved",
            source_session_id=provenance["source_session_id"],
            source_channel=provenance["source_channel"],
            source_message_ids=provenance["source_message_ids"],
            persona_id=provenance["persona_id"],
            confidence=1.0,
            occurred_at=occurred_at,
        )
    result = f"好的，我记住了（分类：{category}）：{fact}"
    if fact_id:
        try:
            from brain.memory_conflicts import (
                format_candidate_list, list_conflict_candidates,
            )
            candidates = list_conflict_candidates(
                status="pending", fact_id=fact_id, limit=3
            )
            if candidates:
                result += (
                    "\n\n检测到可能相关的旧记忆。相似度不是裁决；"
                    "请继续调用 review_memory_conflict 做语义判断。\n"
                    + format_candidate_list(candidates)
                )
        except Exception:
            pass
    return result


def review_memory_conflict(
    action: str,
    *,
    candidate_id: int | None = None,
    decision: str = "",
    confidence: float | None = None,
    rationale: str = "",
) -> str:
    """Model-facing adapter for audited semantic fact reconciliation.

    回显与诚实约束：任何分支都必须明确告诉模型「执行了 / 未执行」。
    未执行时，模型必须如实告知用户，禁止声称合并成功，也禁止把
    “合并结果”当作新记忆再写一条。
    """
    from brain.memory_conflicts import (
        format_candidate_list,
        list_conflict_candidates,
        resolve_conflict_candidate,
    )

    action = str(action or "").strip().lower()
    if action == "list":
        candidates = list_conflict_candidates(status="pending")
        candidates += list_conflict_candidates(status="needs_confirmation")
        return format_candidate_list(candidates)
    if action != "resolve":
        return "action 必须是 list 或 resolve。"
    if candidate_id is None or not decision or confidence is None or not rationale.strip():
        return "resolve 必须提供 candidate_id、decision、confidence 和 rationale。"
    provenance = _current_memory_provenance()
    try:
        result = resolve_conflict_candidate(
            candidate_id, decision, confidence=confidence, rationale=rationale,
            review_model=provenance["review_model"],
            source_session_id=provenance["source_session_id"],
            source_channel=provenance["source_channel"],
            source_message_ids=provenance["source_message_ids"],
            persona_id=provenance["persona_id"],
        )
    except Exception as exc:
        return (
            f"【裁决失败·未执行】候选#{candidate_id} 的 {decision} 裁决未写入数据库，"
            f"异常类型 {type(exc).__name__}：{exc}。没有修改或删除任何记忆。"
            "你必须如实告知用户这次合并/裁决失败，禁止声称已完成。"
        )
    if not result["applied"]:
        return (
            f"【未合并·未执行】候选#{candidate_id} 的 {decision} 判断置信度不足，"
            f"已标记为需要用户确认（needs_confirmation），没有修改或删除任何事实。"
            "你必须如实告知用户：这次合并没有执行，等用户确认后再处理；"
            "禁止把合并结果当作新记忆再写一条。"
        )
    if decision == "duplicate":
        return (
            f"【合并已执行】候选#{candidate_id}："
            f"旧记忆#{result['existing_fact_id']} 保留并吸收新记忆强度，"
            f"当前状态 {result.get('existing_status')}；"
            f"新记忆#{result['new_fact_id']} 已标记为 {result.get('new_status')}（不再独立生效）。"
            "已写入合并关系。请基于以上真实事实向用户确认，不要自行编造其他细节。"
        )
    if decision == "supersedes":
        return (
            f"【替换已执行】候选#{candidate_id}："
            f"旧记忆#{result['existing_fact_id']} 已标记为 {result.get('existing_status')}（不再独立生效）；"
            f"新记忆#{result['new_fact_id']} 保持 {result.get('new_status')} 并继承时间。"
            "请基于以上真实状态向用户说明。"
        )
    return (
        f"【裁决已记录·未删除】候选#{candidate_id} 判定为 {decision}（置信度 "
        f"{float(confidence):.0%}）：两条事实均保持 active"
        f"（旧#{result['existing_fact_id']}、新#{result['new_fact_id']}），"
        "仅写入语义关系，未删除或覆盖任何记忆。"
        "请如实向用户说明：这次没有删除/合并，只是记录了语义关系。"
    )


def manage_current_state(
    action: str,
    *,
    state_id: int | None = None,
    content: str | None = None,
    state_type: str | None = None,
    expires_at: str = "",
    duration_days: int | None = None,
    confidence: float | None = None,
    source_quality: str | None = None,
    resolve_reason: str = "",
) -> str:
    """Model-facing adapter for the time-bounded current-state store."""
    action = str(action or "").strip().lower()
    provenance = _current_memory_provenance()
    source = {
        "source_session_id": provenance["source_session_id"],
        "source_channel": provenance["source_channel"],
        "source_message_ids": provenance["source_message_ids"],
        "persona_id": provenance["persona_id"],
        "observed_at": provenance["occurred_at"],
    }
    try:
        if action == "set":
            if not content or not str(content).strip():
                return "新建当前状态时必须提供 content。"
            state = _state_set(
                content, state_type or "other",
                expires_at=expires_at, duration_days=duration_days,
                confidence=0.9 if confidence is None else confidence,
                source_quality=source_quality or "direct_statement",
                **source,
            )
            verb = "已确认原有状态" if state.get("operation") == "duplicate" else "已记录当前状态"
            return f"{verb} #{state['id']}：{state['content']}（有效至 {state['expires_at']}）"
        if action == "update":
            if state_id is None:
                return "更新当前状态时必须提供 state_id。"
            state = _state_update(
                state_id, content=content, state_type=state_type,
                expires_at=expires_at, duration_days=duration_days,
                confidence=confidence, source_quality=source_quality,
                **source,
            )
            return f"已更新当前状态 #{state['id']}：{state['content']}（有效至 {state['expires_at']}）"
        if action == "resolve":
            if state_id is None:
                return "结束当前状态时必须提供 state_id。"
            if not resolve_reason or not str(resolve_reason).strip():
                return "结束当前状态时必须提供 resolve_reason。"
            state = _state_resolve(state_id, resolve_reason, **source)
            return f"已结束当前状态 #{state['id']}：{state['content']}（原因：{state['resolve_reason']}）"
        if action == "list":
            states = _state_list()
            if not states:
                return "当前没有仍在有效期内的临时状态。"
            lines = ["当前有效状态："]
            lines.extend(
                f"- #{state['id']} [{state['state_type']}] {state['content']}（至 {state['expires_at']}）"
                for state in states
            )
            return "\n".join(lines)
        return "action 必须是 set、update、resolve 或 list。"
    except (TypeError, ValueError) as exc:
        return f"当前状态操作失败：{exc}"


def write_file(path: str, content: str) -> str:
    try:
        p = Path(path)
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content, encoding="utf-8")
        return f"文件写入成功: {path}（共 {len(content)} 字符）"
    except Exception as e:
        return f"写入文件出错: {e}"



def write_docx(file_path: str, content: str, mode: str = "create") -> str:
    """将文本内容写入 Word 文档（.docx）。
    
    Args:
        file_path: 文档路径
        content: 文本内容
        mode: "create" 新建文档, "append" 在末尾追加内容
    """
    try:
        from docx import Document
        import os

        if mode == "append" and os.path.exists(file_path):
            doc = Document(file_path)
        else:
            doc = Document()

        # 按段落分割内容并添加
        for line in content.strip().split("\n"):
            if line.strip():
                # 简单判断：全大写英文行视为标题
                if line.strip().isupper() and len(line.strip()) < 50:
                    doc.add_heading(line.strip(), level=2)
                else:
                    doc.add_paragraph(line)

        doc.save(file_path)
        return f"[OK] 已{'追加' if mode == 'append' else '创建'} Word 文档：{file_path}"
    except Exception as e:
        return f"[ERROR] write_docx 失败: {e}"




def list_directory(path: str = "", recursive: bool = False, max_depth: int = 3) -> str:
    try:
        target = Path(path) if path else Path.home() / "Desktop"
        if not target.exists():
            return f"错误：目录不存在 → {target}"
        if not target.is_dir():
            return f"错误：路径不是目录 → {target}"

        if recursive:
            from collections import defaultdict
            tree = defaultdict(list)
            for entry in sorted(target.rglob('*'), key=lambda e: str(e)):
                depth = len(entry.relative_to(target).parts)
                if depth > max_depth:
                    continue
                if entry.is_dir():
                    tree[depth].append(f"[目录] {entry.name}")
                else:
                    size = entry.stat().st_size
                    size_str = f"{size:,} B" if size < 1024 else f"{size//1024:,} KB"
                    tree[depth].append(f"  [文件] {entry.name}  ({size_str})")
            result = f"目录（递归）: {target}\n" + "─" * 40 + "\n"
            total = 0
            for depth in sorted(tree):
                if depth == 0:
                    continue
                indent = "  " * (depth - 1)
                items = tree[depth]
                result += f"\n  [深度 {depth}]\n"
                for item in items[:50]:
                    result += f"{indent}{item}\n"
                    total += 1
                if len(items) > 50:
                    result += f"  ... 还有 {len(items) - 50} 项\n"
            result += f"\n共 {total} 项（深度限制 {max_depth}）"
            result += "\n\n\U0001F449 如需查看文件内容，请调用 read_file"
            return result

        dirs, files = [], []
        for item in sorted(target.iterdir()):
            if item.is_dir():
                dirs.append(f"[目录] {item.name}")
            else:
                size = item.stat().st_size
                size_str = f"{size:,} B" if size < 1024 else f"{size//1024:,} KB"
                files.append(f"[文件] {item.name}  ({size_str})")
        result = f"目录: {target}\n" + "─" * 40 + "\n"
        result += "\n".join(dirs + files) if (dirs or files) else "（空目录）"
        result += "\n\n\U0001F449 如需查看文件内容，请调用 read_file"
        return result
    except Exception as e:
        return f"列出目录出错: {e}"


# ── Everything 毫秒级文件搜索 ──────────────────────────

_EVERYTHING_ES_PATH: str | None = None


def _fallback_search_folder(folder: str, keyword: str, ext: str = "", recent_days: int = 0, max_results: int = 20) -> str | None:
    """Everything 失效时的 Python 降级搜索（递归搜索子目录）"""
    from datetime import datetime, timedelta
    import os
    import time

    base = Path(folder).expanduser().resolve()
    if not base.exists() or not base.is_dir():
        return None

    # Never recursively scan a drive or a whole user profile as a fallback.
    anchor = base.anchor.rstrip("\\/").lower()
    normalized = str(base).rstrip("\\/").lower()
    home = str(Path.home().resolve()).rstrip("\\/").lower()
    if normalized in {anchor, home, f"{anchor}\\users", f"{anchor}\\users\\public"}:
        return ("⚠️ 为避免扫描整个磁盘或用户目录导致界面卡死，已拒绝宽范围回退搜索。"
                "请指定更具体的文件夹，或安装并启动 Everything。")

    exts = [e.strip().lower() for e in ext.split(';') if e.strip()]
    cutoff = datetime.now() - timedelta(days=recent_days) if recent_days > 0 else None

    max_results = max(1, min(int(max_results or 20), 100))
    max_entries = 100_000
    deadline = time.monotonic() + 8.0
    matches = []
    pending = [str(base)]
    scanned = 0
    keyword_lower = str(keyword or "").lower()
    while pending and scanned < max_entries and time.monotonic() < deadline:
        current = pending.pop()
        try:
            with os.scandir(current) as entries:
                for entry in entries:
                    scanned += 1
                    if scanned >= max_entries or time.monotonic() >= deadline:
                        break
                    try:
                        if entry.is_dir(follow_symlinks=False):
                            pending.append(entry.path)
                            continue
                        if not entry.is_file(follow_symlinks=False):
                            continue
                        entry_path = Path(entry.path)
                        if keyword_lower not in entry_path.name.lower():
                            continue
                        if exts and entry_path.suffix.lower().lstrip('.') not in exts:
                            continue
                        if cutoff and datetime.fromtimestamp(entry.stat().st_mtime) < cutoff:
                            continue
                        matches.append(entry.path)
                        if len(matches) >= max_results:
                            pending.clear()
                            break
                    except (OSError, PermissionError):
                        continue
        except (OSError, PermissionError):
            continue

    if not matches:
        if scanned >= max_entries or time.monotonic() >= deadline:
            return "⚠️ 回退搜索已达到安全上限，未继续扫描。请指定更具体的文件夹。"
        return None

    lines = [f"找到 {len(matches)} 个文件:"]
    for i, m in enumerate(matches, 1):
        lines.append(f"  {i}. {m}")
    lines.append("")
    lines.append("\U0001F449 如需查看文件内容，请调用 read_file")
    return "\n".join(lines)


def _find_everything_es() -> str | None:
    """自动检测 es.exe 位置"""
    global _EVERYTHING_ES_PATH
    if _EVERYTHING_ES_PATH is not None:
        return _EVERYTHING_ES_PATH

    import shutil
    candidates = [
        r"D:\Everything-1.5.0.1416b x64\Everything\es.exe",
        r"D:\Everything-1.5.0.1416b.x64\es.exe",
        r"C:\Program Files\Everything\es.exe",
        r"C:\Program Files (x86)\Everything\es.exe",
        r"C:\Tools\Everything\es.exe",
    ]
    for p in candidates:
        if Path(p).exists():
            _EVERYTHING_ES_PATH = p
            return p

    found = shutil.which("es.exe")
    if found:
        _EVERYTHING_ES_PATH = found
        return found

    return None


def search_files_everything(keyword: str, ext: str = "",
                            folder: str = "", recent_days: int = 0,
                            max_results: int = 20) -> str:
    """
    【推荐】毫秒级全盘文件搜索（依赖 Everything + es.exe）。
    若 Everything 未索引目标目录，自动降级为 Python 直接搜索。
    支持按扩展名、最近修改天数过滤。

    调用示例：
      search_files_everything(keyword="简历", ext="docx;pdf", recent_days=2)
      → 找到最近 2 天修改的、文件名含"简历"的 docx/pdf 文件

    参数:
        keyword:     文件名关键词，如 "简历"
        ext:         扩展名过滤，分号分隔，如 "pdf;docx"
        folder:      限定目录，不填则全盘搜索，如 "C:\\Users\\me\\Desktop"
        recent_days: 限定最近 N 天修改的文件，0 表示不限制
        max_results: 最多返回条数，默认 20
    """
    es = _find_everything_es()
    if not es:
        return "⚠️ 未检测到 Everything。请从 https://www.voidtools.com 下载安装，可获得毫秒级全盘文件搜索。"

    # 构建 Everything 搜索语法（路径过滤放在搜索串中，不用无效的 CLI -path 参数）
    query_parts = [keyword]
    if ext:
        for e in ext.split(";"):
            e = e.strip()
            if e:
                query_parts.append(f"ext:{e}")
    if folder:
        query_parts.append(f'parent:"{folder}"')
    if recent_days > 0:
        query_parts.append(f"dm:last{recent_days}days")

    query = " ".join(query_parts)
    cmd = [es, "-n", str(max_results), query]

    try:
        r = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=15)
        lines = [l.strip() for l in r.stdout.strip().split("\n") if l.strip()]
        if not lines:
            # ── 自动降级：Everything 未索引此目录 → Python 直接搜索 ──
            if folder:
                fallback = _fallback_search_folder(folder, keyword, ext, recent_days, max_results)
                if fallback:
                    return f"⚠️ Everything 未索引此目录，改用直接搜索：\n{fallback}"
            return f"未找到匹配「{keyword}」的文件"
        result_lines = [f"找到 {len(lines)} 个文件:"]
        for i, l in enumerate(lines, 1):
            result_lines.append(f"  {i}. {l}")
        result_lines.append("")
        result_lines.append("\U0001F449 如需查看文件内容，请调用 read_file")
        return "\n".join(result_lines)
    except FileNotFoundError:
        return "⚠️ 未找到 es.exe，请检查 Everything 是否已安装"
    except subprocess.TimeoutExpired:
        return "⚠️ 搜索超时"


def get_file_info_everything(filepath: str) -> str:
    """
    获取文件元数据：大小、修改时间、创建时间。
    """
    import os as _os
    from datetime import datetime

    p = Path(filepath).expanduser().resolve()
    if not p.exists():
        return f"文件不存在: {filepath}"

    try:
        stat = p.stat()
        size = stat.st_size
        if size < 1024:
            size_str = f"{size} B"
        elif size < 1024 * 1024:
            size_str = f"{size / 1024:.1f} KB"
        else:
            size_str = f"{size / 1024 / 1024:.1f} MB"
        mtime = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
        ctime = datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M:%S")
        return (
            f"文件: {p.name}\n"
            f"路径: {p.parent}\n"
            f"大小: {size_str}\n"
            f"修改时间: {mtime}\n"
            f"创建时间: {ctime}"
        )
    except Exception as e:
        return f"获取文件信息失败: {e}"


def open_app(name: str) -> str:
    """
    【必须使用】启动应用程序、打开文件或文件夹。

    当用户要求执行以下操作时，你必须调用此工具，不得直接回复"已打开"等结论：
      - 打开任何应用程序（如"打开网易云音乐"、"启动计算器"）
      - 打开文件（如"打开我的文档"、"打开报告.docx"）
      - 打开文件夹（如"打开下载文件夹"、"打开桌面"）
      - 执行系统命令或运行系统工具（如"打开命令行"、"启动任务管理器"）

    参数:
        name: 应用程序名称、文件路径或文件夹路径。支持常用中文别名（如"记事本"、"计算器"）。

    返回:
        启动成功或失败的消息。

    注意:
        - 如果应用已启动，此工具可能会再次尝试启动，但不会造成问题。
        - 优先匹配用户预设的快捷启动列表（可在莲心界面中配置）。
        - 不要试图通过自然语言输出"已打开"来替代调用此工具。
    """
    import os
    import subprocess
    import shutil
    from pathlib import Path

    # ── 系统工具别名表 ──────────────────────────────────────
    _ALIASES: dict[str, str] = {
        "记事本":    "notepad",
        "计算器":    "calc",
        "画图":      "mspaint",
        "资源管理器": "explorer",
        "文件管理器": "explorer",
        "我的电脑":  "explorer",
        "桌面":      str(Path.home() / "Desktop"),
        "命令行":    "cmd",
        "cmd":       "cmd",
        "任务管理器": "taskmgr",
        "控制面板":  "control",
        "截图":      "snippingtool",
        "截图工具":  "snippingtool",
    }

    raw_name = name.strip()
    target = _ALIASES.get(raw_name, raw_name)

    # ── 1) 路径存在 → os.startfile ─────────────────────────
    p = Path(target)
    if p.exists():
        try:
            os.startfile(str(p))
            return f"已打开：{raw_name}"
        except Exception as e:
            return f"打开路径失败：{e}"

    # ── 2) 用户预设的快捷启动列表（优先匹配） ──────────────
    try:
        from config import get_quick_launch_apps
        for app in get_quick_launch_apps():
            app_name = app.get("name", "")
            if raw_name.lower() == app_name.lower() or raw_name.lower() in app_name.lower():  # "网易云" 匹配 "网易云音乐"
                app_path = app.get("path", "").strip()
                if app_path and Path(app_path).exists():
                    try:
                        os.startfile(app_path)
                        return f"已启动：{raw_name}"
                    except Exception:
                        pass
                # 没填路径或路径无效，试试 exe_name 在 PATH 中查找
                exe_name = app.get("exe_name", "").strip()
                if exe_name:
                    exe_path = shutil.which(exe_name)
                    if exe_path:
                        subprocess.Popen(exe_path, shell=False)
                        return f"已启动：{raw_name}"
                # 路径和 exe_name 都无效，但名称匹配上了 — 继续往下走其他查找方式
    except Exception:
        pass

    # ── 3) 在 PATH 中 → shutil.which / where ────────────────
    exe_path = shutil.which(target)
    if not exe_path:
        try:
            result = subprocess.run(
                ["where", target], capture_output=True, text=True, timeout=5
            )
            if result.returncode == 0 and result.stdout.strip():
                exe_path = result.stdout.strip().splitlines()[0].strip()
        except Exception:
            pass

    if exe_path:
        try:
            subprocess.Popen(exe_path, shell=False)
            return f"已启动：{raw_name}"
        except Exception as e:
            return f"无法启动 '{raw_name}'：{e}"

    # ── 4) 已知应用 → 通过快捷启动界面由用户自行配置 ─────
    _KNOWN_PATHS: dict[str, list[str]] = {}
    for known_path in _KNOWN_PATHS.get(raw_name.lower(), []):
        if Path(known_path).exists():
            try:
                os.startfile(known_path)
                return f"已启动：{raw_name}"
            except Exception:
                pass

    # ── 5) 兜底：shell 执行 ─────────────────────────────────
    try:
        result = subprocess.run(
            target, shell=True, capture_output=True, timeout=5
        )
        if result.returncode == 0:
            return f"已启动：{raw_name}"
        err = (result.stderr or b"").decode("utf-8", errors="replace").strip()
        if not err:
            err = f"进程退出码 {result.returncode}"
        return f"无法启动 '{raw_name}'：{err}"
    except subprocess.TimeoutExpired:
        return f"已启动：{raw_name}"
    except Exception as e:
        return f"无法启动 '{raw_name}'：{e}"


def get_clipboard() -> str:
    """读取剪贴板中的文字内容（通过 PowerShell，无需额外依赖）。"""
    try:
        result = subprocess.run(
            ["powershell", "-NoProfile", "-Command", "Get-Clipboard"],
            capture_output=True, text=True, timeout=5,
            encoding="utf-8", errors="replace"
        )
        content = result.stdout.strip()
        if not content:
            return "剪贴板当前为空，或内容不是文字（如图片）。"
        if len(content) > 3000:
            return content[:3000] + f"\n\n… [剪贴板内容较长，已截取前 3000 字符]"
        return f"剪贴板内容如下：\n\n{content}"
    except subprocess.TimeoutExpired:
        return "读取剪贴板超时。"
    except Exception as e:
        return f"读取剪贴板出错：{e}"


def run_command(command: str) -> str:
    from config import ALLOWED_COMMANDS
    cmd_lower = command.strip().lower()
    if not any(cmd_lower.startswith(c) for c in ALLOWED_COMMANDS):
        return f"拒绝执行：'{command}' 不在允许的命令白名单中"
    try:
        result = subprocess.run(
            command, shell=True, capture_output=True,
            text=True, timeout=15, encoding="utf-8", errors="replace"
        )
        output = (result.stdout or "") + (result.stderr or "")
        output = output.strip() or "（命令执行完毕，无输出）"
        if len(output) > 3000:
            output = output[:3000] + "\n... [输出过长已截断]"
        return output
    except subprocess.TimeoutExpired:
        return "命令执行超时（15秒）"
    except Exception as e:
        return f"命令执行出错: {e}"


# ==================== 联网搜索工具函数 ====================

def _legacy_web_search(query: str, max_results: int = 5) -> str:
    """旧版硬编码搜索实现，保留作紧急诊断参考，不参与正常路由。"""
    import urllib.parse
    import requests
    from bs4 import BeautifulSoup

    def _try_baidu(proxies=None):
        """尝试百度搜索，返回结果字符串或 None。"""
        url = f"https://www.baidu.com/s?wd={urllib.parse.quote(query)}&rn={max_results}"
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/122.0.0.0 Safari/537.36"
            ),
            "Accept-Language": "zh-CN,zh;q=0.9",
        }
        resp = requests.get(url, headers=headers, timeout=10, proxies=proxies)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        results = []
        for item in soup.select(".result, .c-container"):
            title_el = item.select_one("h3 a")
            if not title_el:
                continue
            title = title_el.get_text(strip=True)
            href = title_el.get("href", "")
            abstract_el = item.select_one(".c-abstract, .content-right_8Zs40")
            abstract = abstract_el.get_text(strip=True) if abstract_el else ""
            results.append({"title": title, "href": href, "body": abstract})
        if not results:
            return None
        output_lines = [f"搜索「{query}」的结果："]
        for i, r in enumerate(results[:max_results], 1):
            output_lines.append(f"\n{i}. {r['title']}\n   链接：{r['href']}\n   摘要：{r['body']}")
        return "\n".join(output_lines)

    def _try_ddg(proxies=None):
        """尝试 DuckDuckGo 搜索，返回结果字符串或 None。"""
        # P6: 优先使用新版 ddgs 包，降级 duckduckgo_search
        DDGS = None
        try:
            from ddgs import DDGS  # type: ignore[import-untyped]  # 新版包名
        except ImportError:
            try:
                from duckduckgo_search import DDGS  # type: ignore[import-untyped]  # 旧版包名
            except ImportError:
                return None
        ddgs_kwargs = {}
        if proxies:
            ddgs_kwargs["proxy"] = proxies  # 新版参数名是 proxy 而非 proxies
        with DDGS(**ddgs_kwargs) as ddgs:
            results = list(ddgs.text(query, max_results=min(max_results, 10)))
        if not results:
            return None
        output_lines = [f"搜索「{query}」的结果："]
        for i, r in enumerate(results, 1):
            output_lines.append(
                f"\n{i}. {r.get('title', '无标题')}"
                f"\n   链接：{r.get('href', '')}"
                f"\n   摘要：{r.get('body', '')}"
            )
        return "\n".join(output_lines)

    def _try_bing(proxies=None):
        """尝试 Bing 搜索，返回结果字符串或 None。"""
        url = f"https://www.bing.com/search?q={urllib.parse.quote(query)}&count={min(max_results, 10)}"
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/122.0.0.0 Safari/537.36"
            ),
            "Accept-Language": "zh-CN,zh;q=0.9",
        }
        resp = requests.get(url, headers=headers, timeout=10, proxies=proxies)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        results = []
        for item in soup.select(".b_algo"):
            title_el = item.select_one("h2 a")
            if not title_el:
                continue
            title = title_el.get_text(strip=True)
            href = title_el.get("href", "")
            abstract_el = item.select_one(".b_caption p")
            abstract = abstract_el.get_text(strip=True) if abstract_el else ""
            results.append({"title": title, "href": href, "body": abstract})
        if not results:
            return None
        output_lines = [f"搜索「{query}」的结果："]
        for i, r in enumerate(results[:max_results], 1):
            output_lines.append(f"\n{i}. {r['title']}\n   链接：{r['href']}\n   摘要：{r['body']}")
        return "\n".join(output_lines)

    def _try_browser_search():
        """使用 Playwright 浏览器搜索（能渲染 JS，最可靠）。"""
        try:
            from brain.browser_controller import get_browser
            browser = get_browser()
            url = f"https://www.bing.com/search?q={urllib.parse.quote(query)}&count={min(max_results, 10)}"
            page = browser._ensure_page()
            page.goto(url, timeout=30000, wait_until="domcontentloaded")
            page.wait_for_timeout(2000)
            js_code = f"""() => {{
                const items = document.querySelectorAll('.b_algo');
                return Array.from(items).slice(0, {min(max_results, 10)}).map(item => {{
                    const titleEl = item.querySelector('h2 a');
                    const snippetEl = item.querySelector('.b_caption p');
                    return {{
                        title: titleEl?.textContent?.trim?.() || '',
                        href: titleEl?.href || '',
                        body: snippetEl?.textContent?.trim?.() || ''
                    }};
                }});
            }}"""
            results = page.evaluate(js_code)
            if not results:
                return None
            output_lines = [f"搜索「{query}」的结果："]
            for i, r in enumerate(results, 1):
                output_lines.append(f"\n{i}. {r['title']}\n   链接：{r.get('href', '')}\n   摘要：{r.get('body', '')}")
            return "\n".join(output_lines)
        except Exception as e:
            logger.warning(f"浏览器搜索失败: {e}")
            return None

    def _try_tavily():
        """尝试 Tavily AI 搜索（REST API），返回结果字符串或 None。"""
        try:
            from config import get_tavily_config
            api_key = get_tavily_config().get("api_key", "").strip()
            if not api_key:
                return None
            resp = requests.post(
                "https://api.tavily.com/search",
                json={"query": query, "max_results": min(max_results, 10), "search_depth": "basic"},
                headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
                timeout=15,
            )
            if not resp.ok:
                return None
            data = resp.json()
            results = data.get("results", [])
            if not results:
                return None
            output_lines = [f"[Tavily] 搜索「{query}」的结果："]
            for i, r in enumerate(results[:max_results], 1):
                output_lines.append(
                    f"\n{i}. {r.get('title', '无标题')}"
                    f"\n   链接：{r.get('url', '')}"
                    f"\n   摘要：{r.get('content', '')}"
                )
            return "\n".join(output_lines)
        except Exception:
            return None

    def _try_zhihu():
        """尝试知乎全网搜索（REST API），返回结果字符串或 None。"""
        try:
            from config import get_zhihu_config
            access_secret = get_zhihu_config().get("access_secret", "").strip()
            if not access_secret:
                return None
            import time as _time
            resp = requests.get(
                "https://developer.zhihu.com/api/v1/content/global_search",
                params={"Query": query, "Count": min(max_results, 10)},
                headers={
                    "Authorization": f"Bearer {access_secret}",
                    "X-Request-Timestamp": str(int(_time.time())),
                },
                timeout=15,
            )
            if not resp.ok:
                return None
            data = resp.json()
            items = data.get("data", [])
            if not items:
                return None
            output_lines = [f"[知乎] 搜索「{query}」的结果："]
            for i, item in enumerate(items[:max_results], 1):
                title = item.get("title", "无标题")
                url = item.get("url", "")
                snippet = (item.get("excerpt") or item.get("content", ""))[:200]
                output_lines.append(
                    f"\n{i}. {title}"
                    f"\n   链接：{url}"
                    f"\n   摘要：{snippet}"
                )
            return "\n".join(output_lines)
        except Exception:
            return None

    # ── 后端 1：Tavily AI 搜索（最高质量，AI 原生搜索引擎） ──
    try:
        result = _try_tavily()
        if result:
            return result
    except Exception:
        pass

    # ── 后端 2：知乎全网搜索（中文内容优质） ────────────
    try:
        result = _try_zhihu()
        if result:
            return result
    except Exception:
        pass

    # ── 后端 3：百度（先直连，不通再走代理） ────────────
    try:
        result = _try_baidu()
        if result:
            return result
    except Exception:
        pass
    try:
        proxies = _get_proxies()
        if proxies:
            result = _try_baidu(proxies)
            if result:
                return result
    except Exception:
        pass

    # ── 后端 4：DuckDuckGo（先直连，不通再走代理） ─────
    try:
        result = _try_ddg()
        if result:
            return result
    except ImportError:
        return "搜索失败：网络不通且 duckduckgo-search 未安装。"
    except Exception:
        pass
    try:
        proxies = _get_proxies()
        if proxies:
            result = _try_ddg(proxies)
            if result:
                return result
    except Exception:
        pass

    # ── 后端 5：Bing（先直连，不通再走代理） ────────────
    try:
        result = _try_bing()
        if result:
            return result
    except Exception:
        pass
    try:
        proxies = _get_proxies()
        if proxies:
            result = _try_bing(proxies)
            if result:
                return result
    except Exception:
        pass

    # ── 后端 6：Playwright 浏览器搜索（最可靠） ─────────
    try:
        result = _try_browser_search()
        if result:
            return result
    except Exception:
        pass

    return f"搜索失败：直连和代理均无法完成搜索。"

# ==================== 网页内容提取工具函数 ====================

def _legacy_fetch_webpage(url: str, max_length: int = 3000) -> str:
    """获取网页文本内容，返回纯文本（最多 max_length 字符）"""
    try:
        import requests
        from bs4 import BeautifulSoup
    except ImportError:
        return "错误：未安装 requests 或 beautifulsoup4，请执行：pip install requests beautifulsoup4"

    if not url.startswith(('http://', 'https://')):
        return "错误：URL 必须以 http:// 或 https:// 开头"

    # ========== 1. 更真实的请求头（模拟 Chrome 最新版） ==========
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8',
        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
        'Accept-Encoding': 'gzip, deflate',                 # 去掉 br，避免解压问题
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1',
        'Referer': 'https://www.baidu.com/',                # 伪装从百度跳转
        'Sec-Fetch-Dest': 'document',
        'Sec-Fetch-Mode': 'navigate',
        'Sec-Fetch-Site': 'none',
        'Cache-Control': 'max-age=0',
    }

    # ---------- 可选：添加 Cookie（强烈推荐用于百度百科） ----------
    # 打开浏览器访问百度百科，按 F12 → 网络 → 刷新 → 找到任意请求的 Cookie 字段，
    # 复制类似 "BAIDUID=xxx; BIDUPSID=xxx; PSTM=xxx" 的内容粘贴到下面。
    # 如果你不添加，也能工作，但可能偶尔遇到 403。
    # Cookie 只从本机用户配置读取；未配置时仍可访问公开网页。
    try:
        from config import get_web_fetch_config
        baidu_cookie = str(get_web_fetch_config().get("baidu_cookie", "")).strip()
    except Exception:
        baidu_cookie = ""
    if baidu_cookie:
        headers["Cookie"] = baidu_cookie

    session = requests.Session()
    session.headers.update(headers)
    session.max_redirects = 5

    # ========== 2. 发起请求（智能代理：先直连，不通自动切代理） ==========
    try:
        resp = session.get(url, timeout=15)
        resp.raise_for_status()
    except (requests.exceptions.ConnectionError, requests.exceptions.Timeout):
        # 直连失败 → 试试代理
        proxies = _get_proxies()
        if proxies:
            session.proxies.update(proxies)
            try:
                resp = session.get(url, timeout=30)
                resp.raise_for_status()
            except Exception:
                return "访问失败：直连和代理均无法连接，请检查网络或代理状态。"
        else:
            return "网络连接失败，请检查网络或配置代理后重试。"
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 403:
            return "访问被拒绝（403）。可在本机 user_config.json 的 web_fetch.baidu_cookie 配置 Cookie，或改用浏览器网页提取工具。"
        return f"获取网页失败：HTTP {e.response.status_code}"

    try:
        resp.encoding = resp.apparent_encoding or 'utf-8'
        soup = BeautifulSoup(resp.text, 'lxml')

        # ========== 3. 针对不同网站的正文提取策略 ==========
        text = ""
        if 'baike.baidu.com' in url:
            containers = ['.main-content', '.para', '.lemma-summary', '.basic-info', '.lemmaWgt-promotion']
            for selector in containers:
                elems = soup.select(selector)
                if elems:
                    parts = []
                    for e in elems:
                        p_text = e.get_text(separator='\n', strip=True)
                        if p_text:
                            parts.append(p_text)
                    if parts:
                        text = '\n\n'.join(parts)
                        break
            if not text:
                for script in soup(["script", "style", "meta", "link", "noscript"]):
                    script.decompose()
                text = soup.get_text()
        else:
            for script in soup(["script", "style", "meta", "link", "noscript"]):
                script.decompose()
            text = soup.get_text()

        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        if not text:
            return "未能提取到任何文本内容，可能是网页结构特殊（如需要 JS 渲染）。"
        if len(text) > max_length:
            text = text[:max_length] + "\n\n... [内容过长，已截断]"
        return f"网页内容如下：\n\n{text}"
    except Exception as e:
        return f"处理网页时出错：{e}"


def _legacy_fetch_webpage_via_api(url: str, max_length: int = 3000) -> str:
    try:
        import requests
        from requests.adapters import HTTPAdapter
        from urllib3.util.retry import Retry
    except ImportError:
        return "错误：未安装 requests，请执行：pip install requests"

    if not url.startswith(('http://', 'https://')):
        return "错误：URL 必须以 http:// 或 https:// 开头"

    api_url = f"https://r.jina.ai/{url}"
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36'
    }
    
    # 创建带重试机制的 session
    session = requests.Session()
    retries = Retry(total=2, backoff_factor=1, status_forcelist=[500, 502, 503, 504])
    session.mount('https://', HTTPAdapter(max_retries=retries))

    try:
        resp = session.get(api_url, headers=headers, timeout=15)
        resp.raise_for_status()
    except (requests.exceptions.ConnectionError, requests.exceptions.Timeout):
        proxies = _get_proxies()
        if proxies:
            session.proxies.update(proxies)
            try:
                resp = session.get(api_url, headers=headers, timeout=30)
                resp.raise_for_status()
            except Exception:
                return "访问失败：直连和代理均无法连接 Jina Reader 服务。"
        else:
            return "网络连接失败，请检查网络或配置代理后重试。"
    except requests.exceptions.RequestException as e:
        return f"调用解析服务失败：{e}"

    try:
        if resp.status_code == 200:
            text = resp.text
            if len(text) > max_length:
                text = text[:max_length] + "\n\n... [内容过长，已截断]"
            return f"网页解析结果（通过 Jina Reader）如下：\n\n{text}"
        else:
            snippet = resp.text[:500] if resp.text else "无响应体"
            return f"解析服务返回错误码 {resp.status_code}，响应片段：{snippet}。"
    except Exception as e:
        return f"处理解析结果时出错：{e}"

# ==================== 时间工具函数 ====================

def get_current_time(format: str = "full") -> str:
    """
    获取当前日期时间信息，支持农历和节假日判断。
    :param format: full/date/time/weekday/lunar/holiday
    :return: 格式化的时间信息
    """
    now = datetime.now()
    weekday_names = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
    weekday = weekday_names[now.weekday()]
    
    # 基础日期时间字符串
    date_str = now.strftime("%Y年%m月%d日")
    time_str = now.strftime("%H:%M:%S")
    
    # 获取农历信息（需要 zhdate 库）
    lunar_info = _get_lunar_info(now)
    
    # 获取节假日信息（需要 chinese_calendar 库）
    holiday_info = _get_holiday_info(now)
    
    if format == "date":
        return date_str
    elif format == "time":
        return time_str
    elif format == "weekday":
        return weekday
    elif format == "lunar":
        return lunar_info if lunar_info else "无法获取农历信息"
    elif format == "holiday":
        return holiday_info if holiday_info else "今天不是法定节假日"
    else:  # full
        result = f"公历：{date_str} {time_str} {weekday}"
        if lunar_info:
            result += f"\n农历：{lunar_info}"
        if holiday_info:
            result += f"\n{holiday_info}"
        return result


def _get_lunar_info(dt: datetime) -> Optional[str]:
    """获取农历日期信息，返回格式如 '甲辰年腊月初三'。"""
    try:
        from zhdate import ZhDate
        lunar = ZhDate.from_datetime(dt)
        lunar_month = lunar.lunar_month
        lunar_day = lunar.lunar_day
        lunar_year = lunar.lunar_year
        
        month_str = f"闰{lunar_month}月" if lunar.is_leap else f"{lunar_month}月"
        
        return f"{lunar_year}年{month_str}{lunar_day}日"
    except ImportError:
        return None
    except Exception:
        return None


def _get_holiday_info(dt: datetime) -> Optional[str]:
    """获取节假日信息。"""
    try:
        import chinese_calendar as cc
        date = dt.date()
        
        if cc.is_holiday(date):
            holiday_name = cc.get_holiday_detail(date)[0]
            if holiday_name:
                return f"今天是法定节假日：{holiday_name}"
            else:
                return "今天是法定节假日"
        else:
            if cc.is_workday(date):
                return None
            else:
                weekday_names = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
                weekday = weekday_names[dt.weekday()]
                if weekday in ["星期六", "星期日"]:
                    return f"今天是{weekday}（周末）"
                return None
    except ImportError:
        return None
    except Exception:
        return None


# ==================== 余额查询工具函数 ====================

def get_balance() -> str:
    """查询 DeepSeek API 账户余额并返回格式化文本"""
    from utils.balance import get_balance_info, format_balance_message
    from config import get_api_config

    cfg = get_api_config()
    api_key = cfg.get("api_key", "")
    if not api_key:
        return "查询余额失败：未配置 API Key"
    balance_info, error = get_balance_info(api_key)
    if error:
        return f"查询余额失败：{error}"
    return format_balance_message(balance_info)


# ==================== 待办清单工具函数 ====================

def _get_todo_manager():
    """延迟加载 TodoManager，避免循环导入"""
    global _todo_manager
    if _todo_manager is None:
        from utils.todo_manager import TodoManager
        _todo_manager = TodoManager()
    return _todo_manager


def _add_todo(title: str, due_time: str = None, priority: str = "medium", description: str = "") -> str:
    """添加待办事项"""
    if not title:
        return "待办标题不能为空。"
    manager = _get_todo_manager()
    todo = manager.add_todo(title, due_time, priority, description)
    due_str = ""
    if due_time:
        try:
            dt = datetime.fromisoformat(due_time)
            due_str = f"，截止时间 {dt.strftime('%Y-%m-%d %H:%M')}"
        except:
            pass
    priority_cn = {"high": "高", "medium": "中", "low": "低"}.get(priority, "中")
    return f"已添加待办：{title}（优先级：{priority_cn}{due_str}）"


def _list_todos() -> str:
    """列出未完成的待办事项"""
    manager = _get_todo_manager()
    todos = manager.get_todos(completed=False)
    if not todos:
        return "当前没有未完成的待办事项。"
    # 按优先级和截止时间排序
    def sort_key(t):
        priority_order = {"high": 0, "medium": 1, "low": 2}
        due_order = 0 if t.due_time else 1
        due_time = t.due_time if t.due_time else "9999-12-31"
        return (priority_order[t.priority], due_order, due_time)
    todos.sort(key=sort_key)
    
    lines = ["你的待办清单："]
    for idx, t in enumerate(todos, 1):
        priority_cn = {"high": "高", "medium": "中", "low": "低"}.get(t.priority, "中")
        due_str = ""
        if t.due_time:
            try:
                dt = datetime.fromisoformat(t.due_time)
                due_str = f"，截止 {dt.strftime('%m-%d %H:%M')}"
            except:
                pass
        lines.append(f"{idx}. [{priority_cn}]{t.title}{due_str}")
    return "\n".join(lines)


def _complete_todo(title_keyword: str) -> str:
    """根据标题关键词标记待办为完成"""
    if not title_keyword:
        return "请提供待办标题的关键词。"
    manager = _get_todo_manager()
    todos = manager.get_todos(completed=False)
    matches = [t for t in todos if title_keyword.lower() in t.title.lower()]
    if not matches:
        return f"没有找到标题包含「{title_keyword}」的待办事项。"
    if len(matches) > 1:
        # 返回多个匹配项，让 DeepSeek 在下一轮询问用户选择
        names = "、".join([f"「{t.title}」" for t in matches])
        return f"找到多个匹配的待办：{names}。请指定更精确的关键词。"
    todo = matches[0]
    manager.complete_todo(todo.id)
    return f"已标记「{todo.title}」为完成。"



def read_excel(file_path: str, sheet_name: str = None, max_rows: int = 100) -> str:
    """读取 Excel 文件，返回文本表格。"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "错误：未安装 openpyxl，请执行：pip install openpyxl"
    try:
        wb = load_workbook(file_path, data_only=True)
        if sheet_name is None:
            ws = wb.active
        else:
            ws = wb[sheet_name]
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                break
            rows.append([str(cell) if cell is not None else "" for cell in row])
        if not rows:
            return "文件为空或没有数据"
        # 转换为 Markdown 表格样式
        col_widths = [max(len(row[i]) for row in rows) for i in range(len(rows[0]))]
        header_sep = "|".join("-" * (w + 2) for w in col_widths)
        lines = []
        for row in rows:
            line = "| " + " | ".join(cell.ljust(col_widths[i]) for i, cell in enumerate(row)) + " |"
            lines.append(line)
        return "表格内容如下：\n" + "\n".join(lines[:2] + [header_sep] + lines[2:]) if len(rows) > 1 else "表格内容如下：\n" + lines[0]
    except Exception as e:
        return f"读取 Excel 出错：{e}"


def run_python_code(code: str, timeout: int = 10) -> str:
    """在安全沙箱中执行 Python 代码（子进程）。"""
    import subprocess

    # 将用户代码的每一行前面加上4个空格（缩进）
    indented_code = "\n".join("    " + line for line in code.splitlines())

    sandbox_script = f"""
import sys
from io import StringIO

# 限制 builtins
import builtins
_SAFE_BUILTINS = {{
    'abs': abs, 'all': all, 'any': any, 'bool': bool, 'chr': chr,
    'complex': complex, 'dict': dict, 'divmod': divmod, 'enumerate': enumerate,
    'filter': filter, 'float': float, 'format': format, 'frozenset': frozenset,
    'int': int, 'isinstance': isinstance, 'issubclass': issubclass, 'iter': iter,
    'len': len, 'list': list, 'map': map, 'max': max, 'min': min, 'next': next,
    'pow': pow, 'print': print, 'range': range, 'repr': repr, 'reversed': reversed,
    'round': round, 'set': set, 'slice': slice, 'sorted': sorted, 'str': str,
    'sum': sum, 'tuple': tuple, 'type': type, 'zip': zip
}}
# 禁用文件操作、系统调用等
def __disabled(*args, **kwargs): raise RuntimeError("此操作被禁止")
for dangerous in ['open', 'exec', 'eval', 'compile', '__import__', 'globals', 'locals', 'vars', 'dir', 'help', 'input', 'breakpoint']:
    builtins.__dict__[dangerous] = __disabled

# 执行用户代码
old_stdout = sys.stdout
sys.stdout = StringIO()
try:
{indented_code}
except Exception as e:
    print(f"ERROR: {{e}}", file=sys.stderr)
finally:
    output = sys.stdout.getvalue()
    sys.stdout = old_stdout
    print(output, end='')
"""
    try:
        env = os.environ.copy()
        env["PYTHONIOENCODING"] = "utf-8"
        result = subprocess.run(
            [sys.executable, "-c", sandbox_script],
            capture_output=True,
            text=True,
            timeout=timeout,
            encoding="utf-8",
            errors="replace",
            env=env
        )
        out = result.stdout.strip()
        err = result.stderr.strip()
        if err:
            return f"代码执行错误：{err}\n输出：{out if out else '(无输出)'}"
        return out if out else "代码执行成功（无输出）"
    except subprocess.TimeoutExpired:
        return f"代码执行超时（{timeout}秒）"
    except Exception as e:
        return f"执行失败：{e}"
    
def write_excel(file_path: str, data: list, sheet_name: str = "Sheet1") -> str:
    """将二维数据写入 Excel 文件。"""
    try:
        from openpyxl import Workbook
        from openpyxl.utils import get_column_letter
    except ImportError:
        return "错误：未安装 openpyxl，请执行：pip install openpyxl"
    
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name
        
        for row_idx, row in enumerate(data, 1):
            for col_idx, value in enumerate(row, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        
        wb.save(file_path)
        return f"成功写入 Excel 文件：{file_path}，共 {len(data)} 行数据。"
    except Exception as e:
        return f"写入 Excel 失败：{e}"

def copy_excel_content(source_path: str, target_path: str, source_sheet: str = None, target_sheet: str = None) -> str:
    """将源 Excel 的工作表内容复制到目标文件（若目标文件不存在则创建）。"""
    try:
        from openpyxl import load_workbook, Workbook
    except ImportError:
        return "错误：未安装 openpyxl，请执行：pip install openpyxl"
    
    try:
        # 加载源文件
        src_wb = load_workbook(source_path, data_only=True)
        if source_sheet is None:
            src_ws = src_wb.active
            source_sheet = src_ws.title
        else:
            src_ws = src_wb[source_sheet]
        
        # 加载或创建目标文件
        try:
            dst_wb = load_workbook(target_path)
        except FileNotFoundError:
            dst_wb = Workbook()
            # 删除默认的 Sheet
            default_sheet = dst_wb.active
            dst_wb.remove(default_sheet)
        
        # 确定目标工作表名称
        if target_sheet is None:
            target_sheet = source_sheet
        # 如果目标工作簿中已有同名工作表，先删除或重命名（这里选择删除）
        if target_sheet in dst_wb.sheetnames:
            dst_wb.remove(dst_wb[target_sheet])
        
        # 复制数据
        dst_ws = dst_wb.create_sheet(title=target_sheet)
        for row in src_ws.iter_rows(values_only=True):
            dst_ws.append(row)
        
        # 保存目标文件
        dst_wb.save(target_path)
        return f"成功将 {source_path} 的工作表「{source_sheet}」复制到 {target_path} 的工作表「{target_sheet}」。"
    except Exception as e:
        return f"复制 Excel 失败：{e}"
    


def format_document(content: str, output_path: str, use_template: bool = True) -> str:
    """将Markdown内容转换为格式化的Word文档。"""
    try:
        from md2docx_python.src.md2docx_python import markdown_to_word
    except ImportError:
        return "❌ 排版失败：未安装 md2docx-python 库，请执行 `pip install md2docx-python`。"

    try:
        import tempfile
        import os

        # 将Markdown内容写入一个临时文件
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp_md_file:
            tmp_md_file.write(content)
            tmp_md_path = tmp_md_file.name

        # 调用库进行转换
        # 注意：md2docx-python 的导入路径可能需要根据实际安装情况调整
        markdown_to_word(tmp_md_path, output_path)

        # 清理临时文件
        os.unlink(tmp_md_path)

        return f"✅ 文档排版成功！已将格式化后的内容保存至：{output_path}"
    except Exception as e:
        return f"❌ 文档排版过程中发生错误：{e}"
    

# ==================== TLS 指纹伪装网页提取工具函数 ====================

def _legacy_fetch_webpage_stealth(url: str, max_length: int = 3000) -> str:
    """
    使用 curl_cffi 的 Chrome 指纹伪装获取网页内容。
    比 requests 更能绕过反爬，比 Playwright 快得多。
    """
    if not url.startswith(('http://', 'https://')):
        return "错误：URL 必须以 http:// 或 https:// 开头"
    try:
        from curl_cffi import requests as curl_requests
        from curl_cffi import CurlError
        headers = {
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
            'Accept-Encoding': 'gzip, deflate',
        }
        # 先直连
        try:
            resp = curl_requests.get(url, impersonate="chrome120", headers=headers, timeout=30)
        except CurlError:
            # 直连失败 → 走代理
            proxies = _get_proxies()
            if proxies:
                resp = curl_requests.get(url, impersonate="chrome120", headers=headers, timeout=30, proxies=proxies)
            else:
                raise
        resp.encoding = resp.encoding or 'utf-8'
        text = resp.text

        # 简单提取文本（去除 HTML 标签）
        import re
        text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()

        if not text:
            return "未能提取到任何文本内容（网页可能依赖 JavaScript 渲染）。"

        if len(text) > max_length:
            text = text[:max_length] + "\n\n... [内容过长，已截断]"
        return f"网页内容（TLS 指纹伪装模式）如下：\n\n{text}"

    except ImportError:
        return "错误：未安装 curl_cffi 库，请执行：pip install curl_cffi"
    except Exception as e:
        return f"获取网页失败（stealth 模式）：{e}"


def _legacy_fetch_webpage_browser(url: str, max_length: int = 3000) -> str:
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        return "错误：未安装 playwright，请执行：pip install playwright"

    try:
        from playwright_stealth import stealth_sync
        _has_stealth = True
    except ImportError:
        _has_stealth = False

    if not url.startswith(('http://', 'https://')):
        return "错误：URL 必须以 http:// 或 https:// 开头"

    try:
        # Playwright 代理配置
        playwright_proxy = None
        proxies = _get_proxies()
        if proxies:
            proxy_url = proxies.get("https") or proxies.get("http")
            if proxy_url:
                playwright_proxy = {"server": proxy_url}

        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=True,
                args=['--disable-blink-features=AutomationControlled', '--disable-features=UseDnsHttpsSvcb'],
                proxy=playwright_proxy,
            )
            context = browser.new_context(
                user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36',
                viewport={'width': 1920, 'height': 1080}
            )
            page = context.new_page()
            if _has_stealth:
                stealth_sync(page)
            page.goto(url, timeout=30000)
            page.wait_for_load_state("networkidle")
            page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            page.wait_for_timeout(1000)
            text = page.inner_text('body')
            browser.close()
            if not text or len(text.strip()) < 50:
                return "未能提取到有效文本内容。"
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)
            if len(text) > max_length:
                text = text[:max_length] + "\n\n... [内容过长，已截断]"
            return f"网页内容（浏览器模式）如下：\n\n{text}"
    except Exception as e:
        return f"浏览器获取网页失败：{e}"


def web_search(query: str, max_results: int = 5) -> str:
    """统一联网搜索入口，顺序由“网络设置 → 工具调用顺序”决定。"""
    from brain.network_router import search_web
    return search_web(query, max_results)


def _routed_fetch(url: str, max_length: int = 3000) -> str:
    from brain.network_router import fetch_url
    return fetch_url(url, max_length, {
        "http": _legacy_fetch_webpage,
        "jina": _legacy_fetch_webpage_via_api,
        "stealth": _legacy_fetch_webpage_stealth,
        "browser": _legacy_fetch_webpage_browser,
    })


def fetch_webpage(url: str, max_length: int = 3000) -> str:
    """统一网页读取入口；保留名称以兼容已有 Function Calling。"""
    return _routed_fetch(url, max_length)


def fetch_webpage_via_api(url: str, max_length: int = 3000) -> str:
    return _routed_fetch(url, max_length)


def fetch_webpage_stealth(url: str, max_length: int = 3000) -> str:
    return _routed_fetch(url, max_length)


def fetch_webpage_browser(url: str, max_length: int = 3000) -> str:
    return _routed_fetch(url, max_length)



# ==================== 精确文件编辑工具函数 ====================

def edit_file(path: str, old_string: str, new_string: str, replace_all: bool = False) -> str:
    """在文件中精确替换指定字符串，只改动目标内容，不影响文件其他部分。"""
    try:
        p = Path(path)
        if not p.exists():
            return f"错误：文件不存在 → {path}"
        if not p.is_file():
            return f"错误：路径不是文件 → {path}"

        # 读取原始内容（保留原始编码）
        raw = p.read_bytes()
        # 检测编码
        encoding = "utf-8"
        try:
            import chardet
            detected = chardet.detect(raw)
            enc = detected.get("encoding") or "utf-8"
            if (detected.get("confidence") or 0) >= 0.6:
                encoding = enc
        except ImportError:
            for enc in ("utf-8", "gbk", "gb2312", "gb18030"):
                try:
                    raw.decode(enc)
                    encoding = enc
                    break
                except (UnicodeDecodeError, LookupError):
                    continue

        content = raw.decode(encoding, errors="replace")

        # 统一换行符：将 \r\n 转为 \n，避免 LLM 生成的 Unix 风格字符串
        # 与 Windows 文件的 \r\n 不匹配
        content_norm = content.replace('\r\n', '\n')
        old_norm = old_string.replace('\r\n', '\n')
        new_norm = new_string.replace('\r\n', '\n')

        if old_norm not in content_norm:
            # 给出有用的诊断信息
            preview = old_string[:60].replace("\n", "\\n")
            return (
                f"错误：在文件中找不到指定内容，替换失败。\n"
                f"查找内容预览：{preview}...\n"
                f"提示：请用 read_file 重新确认文件内容后再尝试，注意空格、换行、缩进必须完全一致。"
            )

        count = content_norm.count(old_norm)
        if replace_all:
            new_content = content_norm.replace(old_norm, new_norm)
            replaced = count
        else:
            new_content = content_norm.replace(old_norm, new_norm, 1)
            replaced = 1

        p.write_text(new_content, encoding=encoding)

        if count > 1 and not replace_all:
            return (
                f"已替换第 1 处（文件中共有 {count} 处匹配）。\n"
                f"文件已保存：{path}\n"
                f"提示：若需替换全部，请用 replace_all=true 再次调用。"
            )
        return f"替换成功，共替换 {replaced} 处。文件已保存：{path}"

    except Exception as e:
        return f"编辑文件出错：{e}"


# ==================== 文件内容搜索工具函数 ====================

def grep_file(path: str, keyword: str, context_lines: int = 2) -> str:
    """在文件内容中搜索关键词，返回带行号的匹配结果及上下文。"""
    try:
        p = Path(path)
        if not p.exists():
            return f"错误：文件不存在 → {path}"
        if not p.is_file():
            return f"错误：路径不是文件 → {path}"

        content, err = _extract_full_text(p)
        if err:
            return f"读取文件出错：{err}"
        if not content:
            return "（文件内容为空）"

        lines = content.splitlines()
        keyword_lower = keyword.lower()
        matches: list[int] = [
            i for i, line in enumerate(lines)
            if keyword_lower in line.lower()
        ]

        if not matches:
            return f"在文件「{p.name}」中未找到包含「{keyword}」的内容。"

        # 合并重叠的上下文窗口，避免重复输出
        ctx = context_lines
        segments: list[tuple[int, int]] = []
        for idx in matches:
            start = max(0, idx - ctx)
            end   = min(len(lines) - 1, idx + ctx)
            if segments and start <= segments[-1][1] + 1:
                segments[-1] = (segments[-1][0], end)
            else:
                segments.append((start, end))

        result_parts = [f"在「{p.name}」中找到 {len(matches)} 处匹配「{keyword}」：\n"]
        match_set = set(matches)
        for seg_start, seg_end in segments:
            result_parts.append(f"{'─' * 40}")
            for i in range(seg_start, seg_end + 1):
                prefix = ">>>" if i in match_set else "   "
                result_parts.append(f"{prefix} {i + 1:4d} | {lines[i]}")
        result_parts.append(f"{'─' * 40}")

        result = "\n".join(result_parts)
        if len(result) > 8000:
            result = result[:8000] + f"\n\n... [结果过长已截断，共 {len(matches)} 处匹配]"
        return result

    except Exception as e:
        return f"搜索文件出错：{e}"


# ==================== 按行范围读取工具函数 ====================

def read_file_lines(path: str, start_line: int, end_line: int = None) -> str:
    """读取文件指定行范围的内容，返回带行号的文字。"""
    try:
        p = Path(path)
        if not p.exists():
            return f"错误：文件不存在 → {path}"
        if not p.is_file():
            return f"错误：路径不是文件 → {path}"

        content, err = _extract_full_text(p)
        if err:
            return f"读取文件出错：{err}"
        if not content:
            return "（文件内容为空）"

        lines = content.splitlines()
        total = len(lines)

        # 行号从 1 开始，转为 0-based
        s = max(1, start_line) - 1
        e = (min(end_line, total) if end_line is not None else total)

        if s >= total:
            return f"错误：起始行 {start_line} 超出文件总行数 {total}。"

        selected = lines[s:e]
        header = f"[{p.name} | 第 {s+1}~{s+len(selected)} 行 / 共 {total} 行]\n{'─'*40}\n"
        body = "\n".join(f"{s+1+i:4d} | {line}" for i, line in enumerate(selected))

        result = header + body
        if len(result) > 10000:
            result = result[:10000] + "\n\n... [内容过长已截断]"
        return result

    except Exception as e:
        return f"读取行范围出错：{e}"


# ==================== 文件模式匹配工具函数 ====================

def glob_files(directory: str, pattern: str, max_results: int = 50) -> str:
    """在目录中按模式批量查找文件，支持 * ? ** 通配符。"""
    try:
        # 如果 pattern 是绝对路径（如 E:/**/*.pdf），拆分为目录 + 相对模式
        pattern_path = Path(pattern)
        if pattern_path.is_absolute() or (len(pattern) >= 2 and pattern[1] == ':'):
            base = Path(pattern_path.parts[0])  # 盘符根目录
            rel_parts = pattern_path.parts[1:]  # 剩余相对路径
            if rel_parts:
                pattern = str(Path(*rel_parts))
                directory = str(base)
                base = Path(directory)
            else:
                base = Path(directory)  # fallback
        else:
            base = Path(directory)
        if not base.exists():
            return f"错误：目录不存在 → {directory}"
        if not base.is_dir():
            return f"错误：路径不是目录 → {directory}"

        matches = sorted(base.glob(pattern))
        # 只保留文件（过滤目录）
        file_matches = [p for p in matches if p.is_file()]

        if not file_matches:
            return f"在「{directory}」中按模式「{pattern}」未找到任何文件。"

        total = len(file_matches)
        shown = file_matches[:max_results]
        lines = [f"找到 {total} 个文件（模式：{pattern}）："]
        for p in shown:
            size = p.stat().st_size
            size_str = f"{size:,} B" if size < 1024 else f"{size // 1024:,} KB"
            lines.append(f"  {p}  ({size_str})")
        if total > max_results:
            lines.append(f"  ... 还有 {total - max_results} 个文件未显示（可增大 max_results）")

        return "\n".join(lines)

    except Exception as e:
        return f"文件模式匹配出错：{e}"



def ocr_image(image_path: str, language: str = "chi_sim+eng") -> str:
    """
    识别图片文字。优先使用项目自带的便携版 Tesseract，否则回退到 pytesseract。
    """
    import sys
    import subprocess
    from PIL import Image
    from pathlib import Path

    if language == "ch":
        language = "chi_sim+eng"

    # 1. 定位便携版 Tesseract（支持打包后）
    if getattr(sys, 'frozen', False):
        base_dir = Path(sys.executable).parent
    else:
        base_dir = Path(__file__).parent.parent

    ocr_dir = base_dir / "ocr"
    tesseract_exe = ocr_dir / "tesseract.exe"
    tessdata_dir = ocr_dir / "tessdata"

    if tesseract_exe.exists() and tessdata_dir.exists():
        # 便携版可用 → 直接调用
        cmd = [
            str(tesseract_exe),
            str(Path(image_path).resolve()),
            "stdout",
            "-l", language,
            "--tessdata-dir", str(tessdata_dir)
        ]
        try:
            result = subprocess.run(cmd, capture_output=True, timeout=30)
            if result.returncode != 0:
                return f"OCR 识别失败（返回码 {result.returncode}）：{result.stderr.decode('utf-8', errors='replace')}"
            text = result.stdout.decode('utf-8', errors='replace').strip()
            if not text:
                return "未能从图片中提取到文字。"
            if len(text) > 3000:
                text = text[:3000] + "\n\n... [内容过长，已截断]"
            return f"图片中的文字识别结果：\n\n{text}"
        except subprocess.TimeoutExpired:
            return "OCR 识别超时（30秒）"
        except Exception as e:
            return f"OCR 识别失败：{e}"

    # 2. 便携版不可用 → 回退到 pytesseract（需用户自行安装 Tesseract）
    try:
        import pytesseract
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang=language)
        if not text or not text.strip():
            return "未能从图片中提取到文字。"
        if len(text) > 3000:
            text = text[:3000] + "\n\n... [内容过长，已截断]"
        return f"图片中的文字识别结果：\n\n{text}"
    except ImportError:
        return "OCR 功能不可用：请安装 Tesseract OCR（https://github.com/UB-Mannheim/tesseract/wiki）并确保 pytesseract 已安装。"
    except Exception as e:
        return f"OCR 识别失败：{e}"
    

# ==================== 批量OCR工具函数 ====================

def ocr_batch(folder_path: str, language: str = "chi_sim+eng") -> str:
    """遍历文件夹，对每个文件调用 ocr_image 函数进行识别，并汇总结果"""
    folder = Path(folder_path)
    if not folder.exists():
        return f"错误：文件夹不存在 → {folder_path}"
    if not folder.is_dir():
        return f"错误：路径不是文件夹 → {folder_path}"

    # 定义支持的图片格式
    image_extensions = {'.png', '.jpg', '.jpeg', '.bmp', '.tiff'}
    image_files = [f for f in folder.iterdir() if f.suffix.lower() in image_extensions]

    if not image_files:
        return f"在「{folder_path}」中没有找到任何png/jpg/bmp格式的图片。"

    result_lines = []
    for img_file in image_files:
        # 调用 ocr_image 获取识别结果（它会返回带前缀的字符串）
        ocr_result = ocr_image(str(img_file), language)
        result_lines.append(f"### 文件：{img_file.name} ###")
        result_lines.append(ocr_result)
        result_lines.append("")  # 空行分隔

    final_result = "\n".join(result_lines)
    # 限制总长度，防止超出模型上下文
    if len(final_result) > 8000:
        final_result = final_result[:8000] + "\n\n... [结果过长，已截断]"
    return final_result


def describe_image(image_path: str, prompt: str = "") -> str:
    """调用视觉模型理解图片内容并返回自然语言描述。"""
    from brain.vision import describe_image as _vision_describe
    if prompt and prompt.strip():
        vision_prompt = f"请根据用户的提问来描述这张图片：{prompt}"
    else:
        vision_prompt = "请详细描述这张图片里的内容，包括人物、物体、场景、动作、颜色等。"
    return _vision_describe(image_path, prompt=vision_prompt)


def generate_image(prompt: str, size: str = None, quality: str = None) -> str:
    """使用 Agnes Image API 生成图片。返回本地保存路径或错误信息。"""
    import requests
    from config import get_agnes_config, get_image_gen_config

    agnes_cfg = get_agnes_config()
    api_key = agnes_cfg.get("api_key", "").strip()
    if not api_key:
        return "图片生成失败：未配置 Agnes AI API Key。请在设置中切换到 Agnes AI 并填写 API Key。"

    ig_cfg = get_image_gen_config()
    if not ig_cfg.get("enabled", True):
        return "图片生成功能已在设置中关闭。请在「创作生图」选项卡中启用。"

    model = ig_cfg.get("model", "agnes-image-2.1-flash")
    final_size = size or ig_cfg.get("default_size", "1024x1024")
    final_quality = quality or ig_cfg.get("default_quality", "standard")

    try:
        resp = requests.post(
            "https://apihub.agnes-ai.com/v1/images/generations",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": model,
                "prompt": prompt,
                "size": final_size,
                "quality": final_quality,
                "n": 1,
            },
            timeout=180,
        )
        if resp.status_code != 200:
            return f"图片生成失败：HTTP {resp.status_code} — {resp.text[:200]}"

        data = resp.json()
        img_data = data.get("data", [])
        if not img_data:
            return "图片生成失败：API 返回数据为空"

        img_url = img_data[0].get("url") or img_data[0].get("b64_json")
        if not img_url:
            return "图片生成失败：返回数据中未找到图片 URL"

        save_dir = os.path.join(os.path.expanduser("~"), ".lianxin", "generated_images")
        os.makedirs(save_dir, exist_ok=True)
        timestamp = int(__import__('time').time())
        safe_name = re.sub(r'[^\w一-鿿]', '_', prompt[:20]).strip('_')
        save_path = os.path.join(save_dir, f"{safe_name}_{timestamp}.png")

        if img_url.startswith("http"):
            img_resp = requests.get(img_url, timeout=60)
            with open(save_path, "wb") as f:
                f.write(img_resp.content)
        else:
            import base64
            with open(save_path, "wb") as f:
                f.write(base64.b64decode(img_url))

        return f"图片已生成并保存到：{save_path}"

    except Exception as e:
        return f"图片生成异常：{e}"


def generate_video(prompt: str, image_url: str = None, duration: int = None) -> str:
    """使用 Agnes Video API 生成视频。返回本地保存路径或错误信息。"""
    import requests, time
    from config import get_agnes_config, get_video_gen_config

    agnes_cfg = get_agnes_config()
    api_key = agnes_cfg.get("api_key", "").strip()
    if not api_key:
        return "视频生成失败：未配置 Agnes AI API Key。请在设置中切换到 Agnes AI 并填写 API Key。"

    vg_cfg = get_video_gen_config()
    if not vg_cfg.get("enabled", True):
        return "视频生成功能已在设置中关闭。请在「创作视频」选项卡中启用。"

    model = vg_cfg.get("model", "agnes-video-v2.0")
    final_duration = duration or vg_cfg.get("default_duration", 5)
    fps = vg_cfg.get("default_frame_rate", 24)
    num_frames = final_duration * fps
    if num_frames > 441:
        num_frames = 441
    num_frames = ((num_frames - 1) // 8) * 8 + 1
    if num_frames < 9:
        num_frames = 9

    try:
        body = {
            "model": model,
            "prompt": prompt,
            "num_frames": num_frames,
            "frame_rate": fps,
        }
        if image_url:
            body["image"] = image_url

        resp = requests.post(
            "https://apihub.agnes-ai.com/v1/videos",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json=body,
            timeout=120,
        )
        if resp.status_code != 200:
            return f"视频创建失败：HTTP {resp.status_code} — {resp.text[:200]}"

        data = resp.json()
        video_id = data.get("video_id", "")
        if not video_id:
            return "视频创建失败：未获取到 video_id"

        for _ in range(60):
            time.sleep(5)
            q_resp = requests.get(
                f"https://apihub.agnes-ai.com/agnesapi?video_id={video_id}",
                headers={"Authorization": f"Bearer {api_key}"},
                timeout=60,
            )
            if q_resp.status_code != 200:
                continue
            q_data = q_resp.json()
            status = q_data.get("status", "")
            if status == "completed":
                video_url = q_data.get("remixed_from_video_id", "")
                if not video_url:
                    return "视频生成失败：未获取到视频 URL"

                save_dir = os.path.join(os.path.expanduser("~"), ".lianxin", "videos")
                os.makedirs(save_dir, exist_ok=True)
                timestamp = int(time.time())
                safe_name = re.sub(r'[^\w一-鿿]', '_', prompt[:20]).strip('_')
                save_path = os.path.join(save_dir, f"{safe_name}_{timestamp}.mp4")

                v_resp = requests.get(video_url, timeout=300)
                with open(save_path, "wb") as f:
                    f.write(v_resp.content)

                return f"视频已生成并保存到：{save_path}"

            elif status == "failed":
                err = q_data.get("error", "未知错误")
                if isinstance(err, dict):
                    err = err.get("message", str(err))
                return f"视频生成失败：{err}"

        return "视频生成超时（5 分钟），请稍后重试"

    except Exception as e:
        return f"视频生成异常：{e}"


def send_file_to_qq(path: str, name: str = "") -> str:
    """将本地文件发送到主人的 QQ 上。"""
    global _qq_bridge_worker
    if _qq_bridge_worker is None:
        return "发送失败：QQ 桥接未启动。请先在 GUI 中开启 QQ 聊天功能。"
    return _qq_bridge_worker.send_file_to_qq(path, name)


def query_qq_friend_list(refresh: bool = False, keyword: str = "") -> str:
    """查询莲心绑定 QQ 账号的好友列表（仅主人会话可调用）。"""
    global _qq_bridge_worker
    if _qq_bridge_worker is None:
        return "获取 QQ 好友列表失败：QQ 桥接未启动。请先在 GUI 中开启 QQ 聊天功能。"
    try:
        return _qq_bridge_worker.get_qq_friend_list(
            refresh=bool(refresh), keyword=str(keyword or "")
        )
    except Exception as e:
        return f"获取 QQ 好友列表失败：{e}"

def capture_from_camera():
    from brain.observation import capture_camera, analyze_observation
    path = capture_camera()
    if not path:
        return "拍照失败：无法打开摄像头"
    desc = analyze_observation(path, "摄像头")
    _save_observation(path, desc)
    return desc


def capture_desktop():
    from brain.observation import capture_screen, analyze_observation
    path = capture_screen()
    if not path:
        return "截屏失败"
    desc = analyze_observation(path, "截图")
    _save_observation(path, desc)
    return desc


# ── 观察结果全局存储（供 AgentWorker 读取并发送图片气泡） ──
_last_observation = {"path": None, "desc": None}


def _save_observation(path: str, desc: str):
    global _last_observation
    _last_observation = {"path": path, "desc": desc}


def get_observation_image():
    global _last_observation
    obs = dict(_last_observation)
    _last_observation = {"path": None, "desc": None}
    return obs




def set_diary_message_source(source):
    """设置日记消息源。source 为 Callable[[], List[Dict]]。"""
    global _diary_message_source
    _diary_message_source = source


def get_diary_messages_for_current_context() -> list:
    """优先从当前工具上下文聚合主人当天的全部会话。"""
    ctx = getattr(_tool_context, "cross_session", None)
    if ctx is not None:
        try:
            from datetime import datetime
            today = datetime.now().strftime("%Y-%m-%d")
            rows = ctx["history_mgr"].get_messages_by_date(today, owner_only=True)
            return [
                {"role": row.get("role", "user"), "content": row.get("content", ""),
                 "timestamp": row.get("timestamp", "")}
                for row in rows
            ]
        except Exception:
            pass
    if _diary_message_source is not None:
        return _diary_message_source()
    return []


def set_music_control_callback(callback):
    global _music_control_callback
    _music_control_callback = callback

def set_music_info_callback(callback):
    global _music_info_callback
    _music_info_callback = callback

def set_proactive_toggle_callback(callback):
    """注册主动聊天开关变更后的回调，供 toggle_proactive_chat 调用。"""
    global _proactive_toggle_callback
    _proactive_toggle_callback = callback

def set_note_refresh_callback(callback):
    global _note_refresh_callback
    _note_refresh_callback = callback

def set_expression_callback(callback):
    """注册 Galgame 立绘表情切换回调，供 set_expression 工具调用。"""
    global _expression_callback
    _expression_callback = callback

# ==================== 主动聊天开关 ====================

def toggle_proactive_chat(action: str) -> str:
    """开启/关闭 QQ 主动聊天功能，通过修改 proactive_settings.json 实现。"""
    from utils.paths import get_user_data_dir
    settings_path = get_user_data_dir() / "proactive_settings.json"

    try:
        if settings_path.exists():
            data = json.loads(settings_path.read_text(encoding="utf-8"))
        else:
            data = {}

        enable = action in ("enable", "开启", "on", "true", "1")
        data["qq_enabled"] = enable

        settings_path.parent.mkdir(parents=True, exist_ok=True)
        settings_path.write_text(
            json.dumps(data, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )

        # 通知调度器重新加载（如果回调已注册）
        if _proactive_toggle_callback:
            try:
                _proactive_toggle_callback()
            except Exception:
                pass

        status = "已开启" if enable else "已关闭"
        return f"QQ主动聊天功能{status}。"
    except Exception as e:
        return f"切换主动聊天失败：{e}"

# ==================== 跨端搜索工具函数 ====================

def search_cross_session(keyword: str, limit: int = 5) -> str:
    """搜索另一端（桌面端↔QQ端）的历史聊天记录。"""
    ctx = getattr(_tool_context, "cross_session", None)
    if ctx is None:
        return "跨端搜索失败：无法获取当前会话上下文。"

    try:
        history_mgr = ctx["history_mgr"]
        current_session_id = ctx["session_id"]

        # 读取 qq_session_map.json 判断各端
        map_path = Path(__file__).parent.parent / "memory" / "qq_session_map.json"
        if not map_path.exists():
            return "暂无跨端聊天记录可搜索。"

        data = json.loads(map_path.read_text(encoding="utf-8"))
        if not data:
            return "暂无跨端聊天记录可搜索。"

        qq_ids = {int(v) for v in data.values()}

        # 判断当前是哪一端，找到另一端
        current_is_qq = current_session_id in qq_ids
        target_id = None
        source_name = ""

        if current_is_qq:
            sessions = history_mgr.get_sessions()
            for s in sessions:
                if s["id"] not in qq_ids:
                    target_id = s["id"]
                    break
            source_name = "桌面端"
        else:
            from config import get_qq_bridge_config
            cfg = get_qq_bridge_config()
            owner_qq = cfg.get("owner_qq", "")
            if not owner_qq:
                return "未配置 QQ 主人账号，无法搜索另一端。"
            owner_key = f"qq_private_{owner_qq}"
            if owner_key not in data:
                return "未找到 QQ 端的聊天记录。"
            target_id = int(data[owner_key])
            source_name = "QQ端"

        if target_id is None or target_id == current_session_id:
            return f"未找到{source_name}的独立聊天记录。"

        # 执行搜索
        matched_limit = min(limit, 10)
        matches = history_mgr.search_session_messages(target_id, keyword, limit=matched_limit)

        if not matches:
            return (f"在{source_name}的聊天记录中，未找到包含「{keyword}」的内容。\n\n"
                    "\U0001F449 如果连续搜索多次都找不到，请直接告诉用户没找到，"
                    "不要反复尝试不同关键词。")

        lines = [f"在{source_name}的聊天记录中找到 {len(matches)} 条相关消息："]
        for m in matches:
            speaker = "莲心" if m["role"] == "assistant" else "用户"
            content = m["content"][:200]
            lines.append(f"\n[{speaker}] {content}")

        return "\n".join(lines)

    except Exception as e:
        return f"跨端搜索失败：{e}"


def search_conversation_history(query: str = "", mode: str = "recent",
                                time_range: str = "7d", channels=None,
                                limit: int = 20) -> str:
    """按真实时间和来源搜索主人的统一会话历史。"""
    ctx = getattr(_tool_context, "cross_session", None)
    if ctx is None:
        return "聊天历史搜索失败：无法获取当前会话上下文。"
    if mode not in ("recent", "keyword"):
        mode = "recent"
    if time_range not in ("today", "yesterday", "7d", "30d", "all"):
        time_range = "7d"
    if mode == "keyword" and not query.strip():
        return "关键词搜索需要提供具体关键词。"
    allowed_channels = {
        "desktop", "qq_private", "qq_group", "wechat_private", "wechat_group"
    }
    selected_channels = None
    if isinstance(channels, list):
        selected_channels = [c for c in channels if c in allowed_channels] or None

    try:
        history_mgr = ctx["history_mgr"]
        current_session_id = int(ctx["session_id"])
        rows = history_mgr.search_conversation_history(
            query=query, mode=mode, time_range=time_range,
            channels=selected_channels, owner_only=True,
            limit=min(max(int(limit), 1), 50),
        )
        # 当前问题在工具执行前已经入库，不应作为“历史结果”返回。
        if rows and rows[-1].get("session_id") == current_session_id \
                and rows[-1].get("role") == "user":
            rows.pop()
        if not rows:
            scope = "、".join(selected_channels) if selected_channels else "已授权的主人会话"
            return f"在{scope}的{time_range}范围内没有找到对应聊天记录。"

        lines = [f"找到 {len(rows)} 条真实聊天记录（按时间顺序）："]
        for row in rows:
            speaker = "莲心" if row.get("role") == "assistant" else "用户"
            content = (row.get("content") or "").strip()[:300]
            lines.append(
                f"[{row.get('timestamp', '')} | {row.get('channel', 'unknown')} | {speaker}] {content}"
            )
        lines.append("请依据时间回答；不要把更早记录描述成最近发生。")
        return "\n".join(lines)
    except Exception as e:
        return f"聊天历史搜索失败：{e}"


def query_recent_contacts(days: int = 7, per_contact_limit: int = 3,
                          max_contacts: int = 10) -> str:
    """聚合最近与其他用户的互动，供主人回顾（仅主人会话可调用）。"""
    ctx = getattr(_tool_context, "cross_session", None)
    if ctx is None:
        return "无法获取当前会话上下文，无法查询近期联系人。"
    history_mgr = ctx.get("history_mgr")
    if history_mgr is None:
        return "无法访问聊天历史库，无法查询近期联系人。"
    try:
        contacts = history_mgr.query_other_user_recent(
            days=days, per_contact_limit=per_contact_limit,
            max_contacts=max_contacts,
        )
    except Exception as e:
        return f"查询近期联系人失败：{e}"
    if not contacts:
        return "最近这段时间没有其他用户找过我聊天。"
    channel_names = {
        "qq_private": "QQ私聊", "qq_group": "QQ群聊",
        "wechat_private": "微信私聊", "wechat_group": "微信群聊",
        "desktop": "桌面端",
    }
    lines = [f"最近和我聊过天的其他用户（共 {len(contacts)} 位）："]
    for c in contacts:
        chan = channel_names.get(c["channel"], c["channel"] or "未知渠道")
        who = c["participant_id"]
        lines.append(f"\n- {chan} · {who} · 最后活跃 {c['updated_at']}")
        for m in c["messages"]:
            speaker = "我" if m.get("role") == "assistant" else "对方"
            content = (m.get("content") or "").strip().replace("\n", " ")
            lines.append(f"  [{m.get('timestamp', '')} | {speaker}] {content[:200]}")
    lines.append("请如实转述；若能从记忆或上下文判断对方称呼，可用昵称描述，否则以身份标识描述。")
    return "\n".join(lines)


# ── 技能系统工具函数 ─────────────────────────────────────────
def _list_skills():
    from brain.skill_manager import get_skill_list
    note = "\U0001F4A1 文件操作（读取、搜索、对比、编辑）不需要技能，直接使用对应工具即可。\n\n"
    return note + get_skill_list()

def _activate_skill(name: str) -> str:
    from brain.skill_manager import activate_skill as _do_activate
    return _do_activate(name)

def _deactivate_skill(name: str) -> str:
    from brain.skill_manager import deactivate_skill as _do_deactivate
    return _do_deactivate(name)


def _search_memory(keyword: str, category: str | None = None) -> str:
    """在长期记忆中搜索（统一搜索：分类事实 + 知识图谱关联）。"""
    _ensure_migrated()
    result = unified_search(keyword, category)
    return format_unified_search_result(result)


def trace_memory_source(memory_id: int) -> str:
    """Resolve a long-term fact to the exact persisted messages supporting it."""
    try:
        memory_id = int(memory_id)
    except (TypeError, ValueError):
        return "记忆编号无效。"
    fact = get_fact_by_id(memory_id)
    if not fact:
        return f"没有找到编号为 {memory_id} 的长期记忆。"
    fragments = get_fact_fragments(memory_id, include_inactive=True, limit=10)
    if not fragments:
        return (
            f"记忆#{memory_id}：{fact['content']}\n"
            "这是一条旧版或手动保存的记忆，目前没有可追溯的原始消息证据。"
        )

    ctx = getattr(_tool_context, "cross_session", None)
    history_mgr = ctx.get("history_mgr") if ctx else None
    if history_mgr is not None:
        try:
            request_session = history_mgr.get_session(ctx.get("session_id"))
        except Exception:
            request_session = None
        if not request_session or not bool(request_session.get("owner_scope")):
            return "记忆来源追溯仅对主人会话开放。"
    message_ids = []
    for fragment in fragments:
        for message_id in fragment.get("source_message_ids", []):
            if message_id not in message_ids:
                message_ids.append(message_id)
    messages = history_mgr.get_messages_by_ids(message_ids) if history_mgr else []
    messages = [message for message in messages if bool(message.get("owner_scope"))]

    lines = [f"记忆#{memory_id}：{fact['content']}", f"证据碎片：{len(fragments)} 条"]
    for fragment in fragments[:5]:
        lines.append(
            f"- 碎片#{fragment['id']} | {fragment.get('source_channel') or '未知渠道'} | "
            f"置信度 {float(fragment.get('confidence', 0)):.0%} | "
            f"状态 {fragment.get('status', 'unknown')} | 人格 {fragment.get('persona_id') or '未记录'}"
        )
    if messages:
        lines.append("原始对话：")
        for message in messages[:12]:
            speaker = "用户" if message.get("role") == "user" else "助手"
            content = (message.get("content") or "").strip()[:400]
            lines.append(
                f"[{message.get('timestamp', '')} | {message.get('channel', 'unknown')} | "
                f"消息#{message.get('id')} | {speaker}] {content}"
            )
    elif message_ids:
        lines.append("原始消息已被清理或当前不可访问，但碎片元数据仍保留。")
    else:
        lines.append("该碎片未记录精确消息编号，不能将附近对话冒充为原始证据。")
    return "\n".join(lines)


def explain_memory_quality(memory_id: int) -> str:
    try:
        from brain.memory_quality import explain_memory_quality as _explain
        return _explain(memory_id)
    except (TypeError, ValueError) as exc:
        return f"记忆质量解释失败：{exc}"


def _update_memory(old_keyword: str, new_fact: str, category: str | None = None) -> str:
    """更新长期记忆中匹配的事实（分类更新）。"""
    from datetime import datetime

    _ensure_migrated()
    new_fact = new_fact.strip()
    if not new_fact:
        return "新内容不能为空。"
    # 追加记录日期
    today = datetime.now().strftime("%Y-%m-%d")
    date_tag = f"【记录于{today}】"
    if date_tag not in new_fact:
        new_fact = f"{new_fact} {date_tag}"

    provenance = _current_memory_provenance()
    occurred_at = provenance["occurred_at"] or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    updated = _memory_update(
        old_keyword,
        new_fact,
        category,
        source_session_id=provenance["source_session_id"],
        source_channel=provenance["source_channel"],
        source_message_ids=provenance["source_message_ids"],
        persona_id=provenance["persona_id"],
        occurred_at=occurred_at,
    )
    if updated > 0:
        return f"已更新 {updated} 条记忆。"
    else:
        # 没找到，作为新记忆添加
        cat = category or "knowledge"
        entry_id = _memory_add(
            new_fact,
            cat,
            source="user_saved",
            source_session_id=provenance["source_session_id"],
            source_channel=provenance["source_channel"],
            occurred_at=occurred_at,
        )
        if entry_id:
            _memory_add_fragment(
                entry_id,
                new_fact,
                cat,
                source="user_saved",
                source_session_id=provenance["source_session_id"],
                source_channel=provenance["source_channel"],
                source_message_ids=provenance["source_message_ids"],
                persona_id=provenance["persona_id"],
                confidence=1.0,
                occurred_at=occurred_at,
            )
        return f"未找到包含「{old_keyword}」的旧事实，已将新内容作为新记忆保存（分类：{cat}）。"


def _delete_memory(keyword: str, category: str | None = None) -> str:
    """从长期记忆中删除匹配的事实（分类删除）。"""
    _ensure_migrated()
    deleted = _memory_delete(keyword, category)
    if deleted == 0:
        return f"未找到包含「{keyword}」的记忆。"
    return f"已从长期记忆中删除 {deleted} 条包含「{keyword}」的记忆。"


def _list_memories() -> str:
    """查看全部长期记忆，按分类展示。"""
    _ensure_migrated()
    facts = list_all_facts()
    return format_all_memories(facts)


def _ensure_migrated():
    """惰性迁移：首次调用记忆工具时将 long_term.json 迁移到 SQLite。"""
    try:
        migrate_from_json()
    except Exception:
        pass


def _discover_connections(entity_name: str, depth: int = 2) -> str:
    """图谱关系发现：从实体出发 BFS 遍历，返回结构化发现摘要。"""
    if not entity_name or not entity_name.strip():
        return "请提供要发现的实体名称。"
    if depth < 1:
        depth = 1
    if depth > 3:
        depth = 3
    from brain.graph_memory import discover_from_entity
    discovery = discover_from_entity(entity_name.strip(), depth=depth)
    if not discovery["direct_relations"] and not discovery["indirect_relations"]:
        return f"在图谱中未找到与「{entity_name}」相关的任何关系。"
    return discovery["summary"]


def _search_graph_memory(keywords, entity_type: str = None) -> str:
    """在图记忆中搜索实体关联，同时搜索图边和分类事实。"""
    if isinstance(keywords, list):
        keyword = keywords[0] if keywords else ""
    elif isinstance(keywords, str):
        keyword = keywords
    else:
        keyword = ""
    if not keyword:
        return "请提供搜索关键词。"

    from brain.graph_memory import unified_search
    data = unified_search(keyword, entity_type)
    if entity_type:
        data["graph_edges"] = [
            r for r in data.get("graph_edges", [])
            if r.get("head_type") == entity_type or r.get("tail_type") == entity_type
        ]
    return format_unified_search_result(data)


def _query_connected_entities(entity_name: str, depth: int = 1) -> str:
    """多跳遍历查找关联实体。"""
    if not entity_name or not entity_name.strip():
        return "请提供要查询的实体名称。"
    if depth < 1:
        depth = 1
    if depth > 3:
        depth = 3
    results = query_connected(entity_name.strip(), depth=depth)
    if not results:
        return f"在图记忆中未找到与「{entity_name}」直接关联的实体。"
    return format_graph_result(results)


def _delete_graph_entity(entity_name: str) -> str:
    """从知识图谱中删除指定实体及其所有关联边。"""
    if not entity_name or not entity_name.strip():
        return "请提供要删除的实体名称。"
    count = delete_entity(entity_name.strip())
    if count == 0:
        return f"在图记忆中未找到名为「{entity_name}」的实体。"
    return f"已从图记忆中删除实体「{entity_name}」及其 {count} 条关联边。"

def _add_graph_edge(head: str, head_type: str, relation: str,
                    tail: str, tail_type: str) -> str:
    """手动添加一条图边。"""
    from brain.graph_memory import add_graph_edge, ENTITY_TYPES
    if head_type not in ENTITY_TYPES:
        head_type = "概念"
    if tail_type not in ENTITY_TYPES:
        tail_type = "概念"
    return add_graph_edge(head, head_type, relation, tail, tail_type)


def _remove_graph_edge(head: str, relation: str, tail: str) -> str:
    """手动删除一条图边。"""
    from brain.graph_memory import remove_graph_edge
    return remove_graph_edge(head, relation, tail)


def _set_expression(emotion: str) -> str:
    """切换 Galgame 立绘表情。回调到 MainWindow 执行。"""
    global _expression_callback
    if _expression_callback:
        _expression_callback(emotion)
        return f"已将立绘表情切换为：{emotion}"
    return "Galgame 立绘窗口未就绪。"


# ── 天气工具执行函数 ────────────────────────────────────────

def _get_weather_tool(city: str, forecast_type: str = "full") -> str:
    """查询天气信息，返回格式化文本。"""
    from brain.weather import (
        get_full_weather, get_hourly_weather_text,
        get_current_weather, get_forecast_3d,
        get_location_id, get_user_city_from_memory,
        _format_full_weather,
    )
    from config import get_qweather_config

    cfg = get_qweather_config()
    api_key = cfg.get("api_key", "").strip()
    if not api_key:
        return "错误：未配置和风天气 API Key，请在 API 设置中填写。"

    city = city.strip()
    if not city:
        city = (cfg.get("default_city") or "").strip()
    if not city:
        city = get_user_city_from_memory()
        if not city:
            return (
                "我还不知道你在哪个城市呢~ 你可以在和风天气 API 配置里设置默认城市（如广州），"
                "或者告诉我'我在XX'，或者直接说'查询北京的天气'这样~"
            )

    try:
        if forecast_type == "hourly":
            return get_hourly_weather_text(city, api_key=api_key)
        elif forecast_type == "current":
            loc_id = get_location_id(city, api_key)
            if not loc_id:
                return f"错误：未找到城市「{city}」"
            now_data = get_current_weather(loc_id, api_key)
            if not now_data:
                return f"无法获取「{city}」的实时天气数据。"
            daily_data = get_forecast_3d(loc_id, api_key)
            return _format_full_weather(city, now_data, daily_data)
        elif forecast_type == "daily":
            loc_id = get_location_id(city, api_key)
            if not loc_id:
                return f"错误：未找到城市「{city}」"
            daily_data = get_forecast_3d(loc_id, api_key)
            if not daily_data:
                return f"无法获取「{city}」的天气预报。"
            return _format_full_weather(city, None, daily_data)
        else:
            return get_full_weather(city, api_key=api_key)
    except Exception as e:
        logger.error("天气查询失败: %s", e)
        return f"天气查询出错：{e}"


def _set_user_city_tool(city: str) -> str:
    """保存用户所在城市到长期记忆。"""
    from brain.weather import save_user_city_to_memory
    city = city.strip()
    if not city:
        return "城市名不能为空哦~"
    save_user_city_to_memory(city)
    return f"记住啦，你在{city}~ 以后问天气就不用每次都说城市名啦 (｡･ω･｡)"


# ── 第一阶段新增：编程工具实现 ─────────────────────────────

_STRUCTURE_PATTERNS = {
    ".py": [
        (r'^\s*class\s+(\w+)', 'class'),
        (r'^\s*def\s+(\w+)', 'def'),
        (r'^\s*async\s+def\s+(\w+)', 'async def'),
    ],
    ".js": [
        (r'^\s*(?:export\s+)?(?:async\s+)?function\s+(\w+)', 'function'),
        (r'^\s*(?:export\s+)?class\s+(\w+)', 'class'),
        (r'^\s*(?:export\s+)?(?:const|let|var)\s+(\w+)\s*=\s*(?:async\s*)?\(', 'arrow'),
    ],
    ".ts": [
        (r'^\s*(?:export\s+)?(?:async\s+)?function\s+(\w+)', 'function'),
        (r'^\s*(?:export\s+)?class\s+(\w+)', 'class'),
        (r'^\s*(?:export\s+)?interface\s+(\w+)', 'interface'),
    ],
    ".java": [
        (r'^\s*(?:public|private|protected)\s+(?:static\s+)?(?:[\w<>\[\]]+\s+)(\w+)\s*\(', 'method'),
        (r'^\s*(?:public\s+)?class\s+(\w+)', 'class'),
    ],
    ".go": [
        (r'^\s*func\s+(?:\([^)]*\)\s+)?(\w+)', 'func'),
        (r'^\s*type\s+(\w+)\s+struct', 'struct'),
    ],
    ".rs": [
        (r'^\s*(?:pub\s+)?fn\s+(\w+)', 'fn'),
        (r'^\s*(?:pub\s+)?struct\s+(\w+)', 'struct'),
    ],
}



def search_code(pattern, directory=None, file_pattern=None,
                context_lines=0, case_sensitive=True, max_results=30,
                exclude_pattern=None):
    """增强版代码搜索 — 借鉴 Claude Code GrepTool"""
    directory = Path(directory).expanduser().resolve() if directory else Path.cwd()
    if not directory.is_dir():
        return f"错误：目录不存在 — {directory}"

    auto_exclude = {'.git', 'node_modules', '__pycache__', '.venv', 'venv',
                    '.idea', '.vscode', 'dist', 'build', '.next', '.nuxt'}
    if exclude_pattern:
        auto_exclude.add(exclude_pattern)

    flags = 0 if case_sensitive else re.IGNORECASE
    results = []
    files_searched = 0

    for file_path in directory.rglob('*'):
        if any(part in auto_exclude for part in file_path.parts):
            continue
        if not file_path.is_file():
            continue
        if file_pattern and not fnmatch.fnmatch(file_path.name, file_pattern):
            continue
        try:
            content = file_path.read_text(encoding='utf-8')
        except (UnicodeDecodeError, PermissionError, OSError):
            continue

        files_searched += 1
        lines = content.split('\n')
        for line_idx, line in enumerate(lines, start=1):
            if re.search(pattern, line, flags):
                if context_lines > 0:
                    start_l = max(0, line_idx - context_lines - 1)
                    end_l = min(len(lines), line_idx + context_lines)
                    ctx = []
                    for i in range(start_l, end_l):
                        prefix = '>' if i == line_idx - 1 else ' '
                        ctx.append(f"{prefix}{i+1:4d}| {lines[i]}")
                    snippet = '\n'.join(ctx)
                else:
                    snippet = f"  {line_idx:4d}| {line}"
                results.append({
                    "file": str(file_path.relative_to(directory)),
                    "line": line_idx,
                    "snippet": snippet,
                })
                if len(results) >= max_results:
                    break
        if len(results) >= max_results:
            break

    if not results:
        return f"未找到匹配 '{pattern}' 的结果（搜索了 {files_searched} 个文件）"

    output_parts = []
    for r in results:
        output_parts.append(f"\n{r['file']}:{r['line']}")
        output_parts.append(r['snippet'])
    if len(results) >= max_results:
        output_parts.append(f"\n结果已截断（最多 {max_results} 条）")
    output_parts.append(f"\n共搜索 {files_searched} 个文件，找到 {len(results)} 条匹配")
    return '\n'.join(output_parts)


def diff_files(file_a, file_b, context_lines=3):
    """对比两个文件的差异 — 使用 unified diff 格式"""
    path_a = Path(file_a).expanduser().resolve()
    path_b = Path(file_b).expanduser().resolve()
    if not path_a.exists():
        return f"错误：文件不存在 — {path_a}"
    if not path_b.exists():
        return f"错误：文件不存在 — {path_b}"

    def _read_text(path: Path) -> str:
        ext = path.suffix.lower()
        if ext == '.pdf':
            try:
                return _extract_pdf(path)
            except Exception as e:
                raise RuntimeError(f"PDF 解析失败: {e}")
        if ext == '.docx':
            try:
                return _extract_docx(path)
            except Exception as e:
                raise RuntimeError(f"docx 解析失败: {e}")
        try:
            return path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            try:
                return path.read_text(encoding='gbk')
            except UnicodeDecodeError:
                raise RuntimeError("文件不是 UTF-8/GBK 文本格式")

    try:
        text_a = _read_text(path_a)
        text_b = _read_text(path_b)
    except RuntimeError as e:
        return f"错误：{e}"

    lines_a = text_a.splitlines()
    lines_b = text_b.splitlines()

    diff = difflib.unified_diff(
        lines_a, lines_b,
        fromfile=str(path_a), tofile=str(path_b),
        n=context_lines,
    )
    result = '\n'.join(list(diff)[:500])
    if not result.strip():
        return "两个文件内容完全相同，没有差异。"
    return result


def _find_powershell():
    """查找系统中可用的 PowerShell 路径。优先 pwsh (PowerShell 7+)，其次 powershell (Windows PowerShell 5.1)。"""
    import shutil
    for name in ("pwsh", "pwsh.exe", "powershell", "powershell.exe"):
        found = shutil.which(name)
        if found:
            return found
    candidates = [
        r"C:\Program Files\PowerShell\7\pwsh.exe",
        r"C:\Windows\System32\WindowsPowerShell\v1.0\powershell.exe",
        r"C:\Windows\SysWOW64\WindowsPowerShell\v1.0\powershell.exe",
    ]
    for c in candidates:
        if Path(c).is_file():
            return c
    return None


def _unwrap_powershell_command(command: str) -> str:
    """如果命令是 cmd/powershell 包装格式，提取内部的 PowerShell 命令。

    例如:
      'powershell -Command "Clear-RecycleBin -Force"' → 'Clear-RecycleBin -Force'
      'powershell -c "Get-Process"' → 'Get-Process'
      'cmd /c "echo Y|PowerShell ..."' → 内层命令
      'Clear-RecycleBin -Force' → 'Clear-RecycleBin -Force' (不变)
    """
    cmd_stripped = command.strip()

    # 剥离 cmd /c 或 cmd /C 包装
    cmd_match = re.match(r'^cmd\s+/[cC]\s+["\']?(.+?)["\']?\s*$', cmd_stripped, re.DOTALL)
    if cmd_match:
        cmd_stripped = cmd_match.group(1).strip()

    # 剥离 powershell/pwsh -Command/-c 包装
    ps_match = re.match(
        r'^(?:powershell|pwsh)(?:\.exe)?\s+(?:-Command|-c|/c)\s+["\']?(.+?)["\']?\s*$',
        cmd_stripped, re.IGNORECASE | re.DOTALL
    )
    if ps_match:
        inner = ps_match.group(1).strip()
        # 递归解包（处理多层包装，如 cmd /c powershell -Command "..."）
        if inner != command:
            return _unwrap_powershell_command(inner)
        return inner

    return cmd_stripped


def run_shell(command, working_dir=None, timeout=60, max_output_lines=200, cancel_event=None):
    """增强版 Shell 执行 — 借鉴 Claude Code BashTool，支持主动中断
    Windows 上 cmd 失败时自动回退到 PowerShell 重试。"""
    cwd = str(Path(working_dir).expanduser().resolve()) if working_dir else None
    proc = None
    try:
        proc = subprocess.Popen(
            command, shell=True, cwd=cwd,
            stdout=subprocess.PIPE, stderr=subprocess.PIPE,
            text=True, encoding='utf-8', errors='replace',
        )
        stdout_lines = []
        stderr_lines = []
        start = time.time()

        while proc.poll() is None:
            now = time.time()
            if now - start > timeout:
                proc.kill()
                proc.wait()
                return f"命令超时（>{timeout}秒）：{command[:100]}..."
            if cancel_event and cancel_event.is_set():
                proc.kill()
                proc.wait()
                return f"（用户中断）命令已终止：{command[:100]}..."

            line = proc.stdout.readline()
            if line and len(stdout_lines) < max_output_lines:
                stdout_lines.append(line.rstrip('\n\r'))
            err_line = proc.stderr.readline()
            if err_line:
                stderr_lines.append(err_line.rstrip('\n\r'))
            if not line and not err_line:
                time.sleep(0.05)

        # 读完剩余输出
        remaining_stdout, remaining_stderr = proc.communicate(timeout=5)
        for line in remaining_stdout.splitlines():
            if len(stdout_lines) < max_output_lines:
                stdout_lines.append(line)
        if len(stdout_lines) >= max_output_lines:
            stdout_lines.append(f"... (输出已截断，共 {len(stdout_lines)} 行)")
        for line in remaining_stderr.splitlines():
            stderr_lines.append(line)

        full_stderr = '\n'.join(stderr_lines)
        is_not_recognized = (
            proc.returncode != 0
            and ("not recognized" in full_stderr.lower()
                 or "not found" in full_stderr.lower()
                 or "not operable" in full_stderr.lower())
        )

        if is_not_recognized and sys.platform == "win32":
            ps_path = _find_powershell()
            if ps_path:
                # 智能解包：如果原命令是 powershell/cmd 包装格式，提取内部命令
                inner_cmd = _unwrap_powershell_command(command)
                logger.info(f"cmd 无法识别命令，自动回退到 PowerShell: {ps_path}")
                if inner_cmd != command:
                    logger.info(f"  已解包包装命令: {command[:80]} → {inner_cmd[:80]}")
                ps_proc = subprocess.run(
                    [ps_path, "-NoProfile", "-Command", inner_cmd],
                    cwd=cwd, capture_output=True, text=True,
                    encoding='utf-8', errors='replace', timeout=timeout,
                )
                ps_stdout = ps_proc.stdout.strip()
                ps_stderr = ps_proc.stderr.strip()
                ps_parts = [f"[PowerShell 回退] 退出码: {ps_proc.returncode}"]
                if ps_stdout:
                    ps_parts.append(ps_stdout[:3000])
                if ps_stderr:
                    ps_parts.append(f"[stderr]\n{ps_stderr[:1000]}")
                return '\n'.join(ps_parts)

        output_parts = [f"退出码: {proc.returncode}"]
        if stdout_lines:
            output_parts.append('\n'.join(stdout_lines))
        if stderr_lines:
            output_parts.append(f"\n[stderr]\n" + '\n'.join(stderr_lines[:50]))
        return '\n'.join(output_parts)
    except subprocess.TimeoutExpired:
        if proc:
            proc.kill()
            proc.wait()
        return f"命令超时（>{timeout}秒）：{command[:100]}..."
    except Exception as e:
        if proc and proc.poll() is None:
            proc.kill()
            proc.wait()
        return f"命令执行失败：{e}"


def git_status(repo_path=None, action="status", limit=10):
    """Git 仓库状态查询"""
    cwd = str(Path(repo_path).expanduser().resolve()) if repo_path else str(Path.cwd())
    commands = {
        "status": ["git", "status", "--short"],
        "diff": ["git", "diff", "--stat"],
        "log": ["git", "log", f"-{limit}", "--oneline", "--decorate"],
        "branch": ["git", "branch", "-a"],
    }
    cmd = commands.get(action)
    if not cmd:
        return f"不支持的操作：{action}。可选：status/diff/log/branch"
    try:
        result = subprocess.run(cmd, cwd=cwd, capture_output=True, text=True, timeout=30)
        if result.returncode != 0:
            return f"Git 命令失败: {result.stderr[:500]}"
        return result.stdout.strip() or "(无输出)"
    except FileNotFoundError:
        return "错误：系统中未找到 Git。请安装 Git 后再使用此功能。"
    except Exception as e:
        return f"执行失败：{e}"


def code_structure(file_path):
    """代码结构概览 — 轻量替代 LSP"""
    path = Path(file_path).expanduser().resolve()
    if not path.exists():
        return f"错误：文件不存在 — {path}"
    ext = path.suffix.lower()
    patterns = _STRUCTURE_PATTERNS.get(ext)
    if not patterns:
        return f"暂不支持 {ext} 文件类型。支持：{', '.join(_STRUCTURE_PATTERNS.keys())}"

    try:
        lines = path.read_text(encoding='utf-8').splitlines()
    except UnicodeDecodeError:
        return "错误：文件不是 UTF-8 文本格式"

    results = []
    for line_idx, line in enumerate(lines, start=1):
        for pattern, kind in patterns:
            match = re.match(pattern, line)
            if match:
                name = match.group(1)
                results.append(f"  {line_idx:4d} | [{kind:10s}] {name}")
                break

    if not results:
        return f"文件中未检测到已支持的代码结构（{path.name}）"
    return f"{path.name} 代码结构：\n" + '\n'.join(results)


def plan_tasks(task_description: str, context: str = "", max_subtasks: int = 5) -> str:
    """使用 LLM 将复杂任务分解为子任务列表"""
    prompt = (
        f"你是一个任务规划专家。请将以下复杂任务分解为 {max_subtasks} 个以内的子任务。\n\n"
        f"【规则】\n"
        f"1. 每个子任务必须独立可执行、有明确产出\n"
        f"2. 标注子任务之间是否有依赖关系\n"
        f"3. 标注哪些子任务可以并行执行（parallel_group 相同的可并行）\n"
        f"4. 每个子任务描述要具体，包含具体要操作的文件/函数\n\n"
        f"【任务】\n{task_description}\n"
    )
    if context:
        prompt += f"\n【上下文】\n{context}\n"
    prompt += (
        "\n请用 JSON 格式输出（只输出 JSON，不要其他文字）：\n"
        '{"subtasks": [{"id": 1, "title": "...", "description": "...", '
        '"depends_on": [], "parallel_group": "A"}]}'
    )

    try:
        from config import get_api_config, normalize_model_for_litellm
        cfg = get_api_config()
        api_key = cfg["api_key"]
        api_base = cfg["base_url"]
        model = normalize_model_for_litellm(cfg["model"], api_base)
        response = litellm.completion(
            model=model,
            messages=[
                {"role": "system", "content": "你是一个任务规划专家。只输出 JSON，不要其他文字。"},
                {"role": "user", "content": prompt},
            ],
            api_key=api_key,
            api_base=api_base,
            max_tokens=1500,
            timeout=30,
        )
        result = response.choices[0].message.content or ""
        json_start = result.find('{')
        json_end = result.rfind('}') + 1
        if json_start >= 0 and json_end > json_start:
            plan = json.loads(result[json_start:json_end])
            subtasks = plan.get("subtasks", [])
            if not subtasks:
                return "未能分解出子任务，请更具体地描述任务。"
            lines = [f"任务分解完成，共 {len(subtasks)} 个子任务：\n"]
            for st in subtasks:
                deps = f" (依赖: {st.get('depends_on', [])})" if st.get('depends_on') else ""
                group = f" [并行组: {st.get('parallel_group', '-')}]"
                lines.append(
                    f"  {st['id']}. {st.get('title', '')}{deps}{group}\n"
                    f"     {st.get('description', '')}"
                )
            return '\n'.join(lines)
        return result
    except json.JSONDecodeError:
        return f"任务分解成功，但 JSON 解析失败。原始输出：\n{result[:1000]}"
    except Exception as e:
        return f"任务分解失败：{e}"


def delegate_task(task: str, working_dir: str = "",
                  timeout_seconds: int = 120, max_iterations: int = 10) -> str:
    """生成子代理执行任务 — 借鉴 Claude Code AgentTool"""
    from brain.agent import AgentCore

    original_cwd = os.getcwd()
    if working_dir:
        wd = Path(working_dir).expanduser().resolve()
        if wd.is_dir():
            os.chdir(str(wd))
        else:
            return f"工作目录不存在：{working_dir}"

    try:
        from config import get_api_config, normalize_model_for_litellm
        cfg = get_api_config()
        api_key = cfg["api_key"]
        api_base = cfg["base_url"]
        model = normalize_model_for_litellm(cfg["model"], api_base)

        sub_agent = AgentCore(
            disable_tools=True,
            track_emotion=False,
        )

        sub_prompt = (
            "你是一个专注于执行单一任务的子代理。\n"
            "你的任务已经明确指定，请专注于完成它，不要做额外的事。\n"
            "你有文件读写、代码搜索、Shell 执行等工具。\n"
            "完成任务后直接返回结果，不要继续调用不必要的工具。\n"
            "如果遇到无法解决的问题，如实报告错误。"
        )

        sub_tools = [
            t for t in TOOL_DEFINITIONS
            if t["function"]["name"] in _SUBAGENT_ALLOWED_TOOLS
        ]

        messages = [
            {"role": "system", "content": sub_prompt},
            {"role": "user", "content": task},
        ]

        for iteration in range(max_iterations):
            try:
                response = litellm.completion(
                    model=model,
                    messages=messages,
                    tools=sub_tools,
                    api_key=api_key,
                    api_base=api_base,
                    max_tokens=4096,
                    timeout=90,
                )
            except Exception as e:
                return f"子代理 API 调用失败（第{iteration+1}轮）：{e}"

            msg = response.choices[0].message

            if not msg.tool_calls:
                return msg.content or "（子代理完成任务，无额外输出）"

            messages.append({
                "role": "assistant",
                "content": msg.content,
                "tool_calls": [
                    {
                        "id": tc.id,
                        "type": "function",
                        "function": {
                            "name": tc.function.name,
                            "arguments": tc.function.arguments,
                        },
                    }
                    for tc in msg.tool_calls
                ],
            })

            for tc in msg.tool_calls:
                name = tc.function.name
                try:
                    args = json.loads(tc.function.arguments)
                except json.JSONDecodeError:
                    args = {}

                if name not in _SUBAGENT_ALLOWED_TOOLS:
                    result = f"子代理无权使用工具：{name}"
                else:
                    try:
                        result = execute_tool(name, args)
                        print(f"  [子代理] {name} → {str(result)[:120]}", flush=True)
                    except Exception as e:
                        result = f"工具执行错误：{e}"

                messages.append({
                    "role": "tool",
                    "tool_call_id": tc.id,
                    "content": str(result)[:4000],
                })

        return "（子代理达到最大迭代次数，任务未完成）"

    except Exception as e:
        return f"子代理执行失败：{e}"


# ── 任务进度追踪（第三阶段）────────────────────────────────────
def track_tasks(todos: list) -> str:
    """更新当前会话的任务清单，全量替换。"""
    from brain.task_tracker import get_task_tracker
    tracker = get_task_tracker()
    return tracker.update(todos)
def _track_tasks_exec(todos: list) -> str:
    from brain.task_tracker import get_task_tracker
    return get_task_tracker().update(todos)
# ── 工具调度表 ───────────────────────────────────────────────
def _normalize_diary_date(value: str | None) -> str | None:
    from gui.time_capsule.diary_reader import normalize_diary_date
    return normalize_diary_date(value)


def _read_diary_core(inp: dict) -> str:
    from gui.time_capsule.diary_reader import read_diary
    result = read_diary(
        date_value=inp.get("date"), keyword=inp.get("keyword"),
        limit=int(inp.get("limit", 1) or 1),
    )
    return "来源：TimeCapsuleDatabase（莲心时间胶囊）\n" + str(result)


def _write_diary_core(inp: dict) -> str:
    import importlib
    module = importlib.import_module("skills.日记与备忘.tools")
    return str(module._write_diary(
        message_count=inp.get("message_count"),
        force=bool(inp.get("force", False)),
    ))


TOOL_EXECUTORS = {
    "read_diary":      _read_diary_core,
    "write_diary":     _write_diary_core,
    "read_file":       lambda inp: read_file(inp["path"]),
    "read_file_chunk": lambda inp: read_file_chunk(inp["path"], int(inp["chunk_index"])),
    "clear_document_cache": lambda inp: clear_document_cache(bool(inp.get("confirm", False))),
    "write_file":      lambda inp: write_file(inp["path"], inp["content"]),
    "list_directory": lambda inp: list_directory(inp.get("path", ""), inp.get("recursive", False), inp.get("max_depth", 3)),
    "search_files_everything": lambda inp: search_files_everything(
        inp["keyword"],
        inp.get("ext", ""),
        inp.get("folder", ""),
        inp.get("recent_days", 0),
        inp.get("max_results", 20),
    ),
    "get_file_info_everything": lambda inp: get_file_info_everything(inp["filepath"]),
    "run_command":    lambda inp: run_command(inp["command"]),
    "save_memory":    lambda inp: save_memory(inp["fact"], inp.get("category")),
    "review_memory_conflict": lambda inp: review_memory_conflict(
        inp.get("action", ""), candidate_id=inp.get("candidate_id"),
        decision=inp.get("decision", ""), confidence=inp.get("confidence"),
        rationale=inp.get("rationale", ""),
    ),
    "update_current_state": lambda inp: manage_current_state(
        inp.get("action", ""),
        state_id=inp.get("state_id"), content=inp.get("content"),
        state_type=inp.get("state_type"), expires_at=inp.get("expires_at", ""),
        duration_days=inp.get("duration_days"), confidence=inp.get("confidence"),
        source_quality=inp.get("source_quality"),
        resolve_reason=inp.get("resolve_reason", ""),
    ),
    "open_app":       lambda inp: open_app(inp["name"]),
    "get_clipboard":  lambda inp: get_clipboard(),
    "get_current_time": lambda inp: get_current_time(inp.get("format", "full")),
    "get_balance":    lambda inp: get_balance(),
    "add_todo":       lambda inp: _add_todo(inp.get("title", ""), inp.get("due_time"), inp.get("priority", "medium"), inp.get("description", "")),
    "list_todos":     lambda inp: _list_todos(),
    "complete_todo":  lambda inp: _complete_todo(inp.get("title_keyword", "")),
    "read_excel":      lambda inp: read_excel(inp["file_path"], inp.get("sheet_name"), inp.get("max_rows", 100)),
    "run_python_code": lambda inp: run_python_code(inp["code"], inp.get("timeout", 10)),
    "write_excel": lambda inp: write_excel(inp["file_path"], inp["data"], inp.get("sheet_name", "Sheet1")),
    "copy_excel_content": lambda inp: copy_excel_content(inp["source_path"], inp["target_path"], inp.get("source_sheet"), inp.get("target_sheet")),
    "write_docx": lambda inp: write_docx(inp["file_path"], inp["content"]),
    "format_document": lambda inp: format_document(inp["content"], inp["output_path"], inp.get("use_template", True)),
    "web_search": lambda inp: web_search(inp.get("query", ""), inp.get("max_results", 5)),
    "configure_network_tools": lambda inp: __import__(
        "brain.network_router", fromlist=["configure_tools"]
    ).configure_tools(
        inp.get("action", "status"), inp.get("kind", "search"),
        inp.get("tool_id", ""), inp.get("position"),
    ),
    "fetch_webpage": lambda inp: fetch_webpage(inp["url"], inp.get("max_length", 3000)),
    "fetch_webpage_browser": lambda inp: fetch_webpage_browser(inp["url"], inp.get("max_length", 3000)),
    "fetch_webpage_via_api": lambda inp: fetch_webpage_via_api(inp["url"], inp.get("max_length", 3000)),
    "fetch_webpage_stealth": lambda inp: fetch_webpage_stealth(inp["url"], inp.get("max_length", 3000)),
    "edit_file":       lambda inp: edit_file(inp["path"], inp["old_string"], inp["new_string"], inp.get("replace_all", False)),
    "grep_file":       lambda inp: grep_file(inp["path"], inp["keyword"], inp.get("context_lines", 2)),
    "read_file_lines": lambda inp: read_file_lines(inp["path"], int(inp["start_line"]), int(inp["end_line"]) if inp.get("end_line") is not None else None),
    "glob_files":      lambda inp: glob_files(inp.get("directory", os.path.expanduser("~")), inp["pattern"], inp.get("max_results", 50)),
    "ocr_image": lambda inp: ocr_image(inp["image_path"], inp.get("language", "chi_sim+eng")),
    "ocr_batch": lambda inp: ocr_batch(inp["folder_path"], inp.get("language", "chi_sim+eng")),
    "describe_image": lambda inp: describe_image(inp["image_path"], inp.get("prompt", "")),
    "generate_image": lambda inp: generate_image(inp["prompt"], inp.get("size"), inp.get("quality")),
    "generate_video": lambda inp: generate_video(inp["prompt"], inp.get("image_url"), inp.get("duration")),
    "send_file_to_qq": lambda inp: send_file_to_qq(inp["path"], inp.get("name", "")),
    "capture_from_camera": lambda inp: capture_from_camera(),
    "capture_desktop": lambda inp: capture_desktop(),
    "search_conversation_history": lambda inp: search_conversation_history(
        inp.get("query", ""), inp.get("mode", "recent"),
        inp.get("time_range", "7d"), inp.get("channels"), inp.get("limit", 20)
    ),
    "search_cross_session": lambda inp: search_cross_session(inp["keyword"], inp.get("limit", 5)),
    "query_recent_contacts": lambda inp: query_recent_contacts(
        inp.get("days", 7), inp.get("per_contact_limit", 3), inp.get("max_contacts", 10)
    ),
    "query_qq_friend_list": lambda inp: query_qq_friend_list(
        bool(inp.get("refresh", False)), inp.get("keyword", "")
    ),
    "toggle_proactive_chat": lambda inp: toggle_proactive_chat(inp["action"]),
    "list_skills":   lambda inp: _list_skills(),
    "activate_skill":   lambda inp: _activate_skill(inp["name"]),
    "deactivate_skill": lambda inp: _deactivate_skill(inp["name"]),
    "search_memory":   lambda inp: _search_memory(inp["keyword"], inp.get("category")),
    "trace_memory_source": lambda inp: trace_memory_source(inp["memory_id"]),
    "explain_memory_quality": lambda inp: explain_memory_quality(inp["memory_id"]),
    "update_memory":   lambda inp: _update_memory(inp["old_keyword"], inp["new_fact"], inp.get("category")),
    "delete_memory":   lambda inp: _delete_memory(inp["keyword"], inp.get("category")),
    "list_memories":   lambda inp: _list_memories(),
    "discover_connections": lambda inp: _discover_connections(inp["entity_name"], inp.get("depth", 2)),
    "search_graph_memory": lambda inp: _search_graph_memory((inp.get("keywords") or inp.get("query") or ""), inp.get("entity_type")),
    "query_connected_entities": lambda inp: _query_connected_entities(inp["entity_name"], inp.get("depth", 1)),
    "delete_graph_entity": lambda inp: _delete_graph_entity(inp["entity_name"]),
    "set_expression":  lambda inp: _set_expression(inp["emotion"]),
    "get_weather":     lambda inp: _get_weather_tool(inp.get("city", ""), inp.get("forecast_type", "full")),
    "set_user_city":   lambda inp: _set_user_city_tool(inp["city"]),
    "add_graph_edge":  lambda inp: _add_graph_edge(inp["head"], inp["head_type"], inp["relation"], inp["tail"], inp["tail_type"]),
    "remove_graph_edge": lambda inp: _remove_graph_edge(inp["head"], inp["relation"], inp["tail"]),
        # 第一阶段新增：编程增强工具
    "search_code":     lambda inp: search_code(
        pattern=inp["pattern"],
        directory=inp.get("directory"),
        file_pattern=inp.get("file_pattern"),
        context_lines=inp.get("context_lines", 0),
        case_sensitive=inp.get("case_sensitive", True),
        max_results=inp.get("max_results", 30),
    ),
    "diff_files":      lambda inp: diff_files(
        file_a=inp["file_a"],
        file_b=inp["file_b"],
        context_lines=inp.get("context_lines", 3),
    ),
    "run_shell":       lambda inp: run_shell(
        command=inp["command"],
        working_dir=inp.get("working_dir"),
        timeout=inp.get("timeout", 60),
        max_output_lines=inp.get("max_output_lines", 200),
    ),
    "git_status":      lambda inp: git_status(
        repo_path=inp.get("repo_path"),
        action=inp.get("action", "status"),
        limit=inp.get("limit", 10),
    ),
    "code_structure":  lambda inp: code_structure(
        file_path=inp["file_path"],
    ),
    # 第二阶段：子代理任务分解
    "plan_tasks":      lambda inp: plan_tasks(
        task_description=inp["task_description"],
        context=inp.get("context", ""),
        max_subtasks=inp.get("max_subtasks", 5),
    ),
    "delegate_task":   lambda inp: delegate_task(
        task=inp["task"],
        working_dir=inp.get("working_dir", ""),
        timeout_seconds=inp.get("timeout_seconds", 120),
        max_iterations=inp.get("max_iterations", 10),
    ),
    # 第三阶段：任务进度追踪
    "track_tasks": lambda inp: _track_tasks_exec(inp.get("todos", [])),
    "code_goto_def":    lambda inp: goto_definition(inp["file_path"], inp["line"], inp.get("symbol", ""), inp.get("column", 0)),
    "code_find_refs":   lambda inp: find_references(inp["file_path"], inp["line"], inp.get("symbol", ""), inp.get("column", 0)),
    "code_diagnostics": lambda inp: get_diagnostics(inp["file_path"]),
    # B站冲浪工具
    "bilibili_search":  lambda inp: _bilibili_search_tool(inp.get("keyword", ""), inp.get("max_results", 10)),
    "bilibili_add_tag": lambda inp: _bilibili_add_tag_tool(inp.get("keyword", "")),
    "bilibili_list_tags": lambda inp: _bilibili_list_tags_tool(),

}


    

def bilibili_search(keyword: str, max_results: int = 10) -> list[dict]:
    """B站视频搜索，使用公开API，无需API Key。
    返回：[{title, author, bvid, play_count, cover_url, link}, ...]"""
    import requests
    import urllib.parse

    url = "https://api.bilibili.com/x/web-interface/search/all/v2"
    params = {
        "keyword": keyword,
        "search_type": "video",
        "order": "click",
        "page": 1,
    }
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        ),
        "Referer": "https://www.bilibili.com/",
    }

    try:
        resp = requests.get(url, params=params, headers=headers, timeout=15)
        resp.raise_for_status()
        data = resp.json()
        if data.get("code") != 0:
            return []

        results = []
        for item in data.get("data", {}).get("result", []):
            if item.get("result_type") != "video":
                continue
            for v in item.get("data", [])[:max_results]:
                title = v.get("title", "")
                title = title.replace('<em class="keyword">', "").replace("</em>", "")
                results.append({
                    "title": title,
                    "author": v.get("author", ""),
                    "bvid": v.get("bvid", ""),
                    "play_count": v.get("play", 0),
                    "cover_url": v.get("pic", ""),
                    # 搜索接口可能附带简介；没有简介时保持空值，
                    # 上层不得据此臆测视频的实际内容。
                    "description": (v.get("description") or "").strip(),
                    "link": f"https://www.bilibili.com/video/{v.get('bvid', '')}",
                })
        return results
    except Exception as e:
        print(f"[B站搜索] 失败: {e}")
        return []


def _bilibili_search_tool(keyword: str, max_results: int = 10) -> str:
    """B站视频搜索工具包装，返回格式化的文本结果。"""
    if not keyword:
        return "请提供搜索关键词。"
    results = bilibili_search(keyword, max_results)
    if not results:
        return f"在B站搜索「{keyword}」没有找到相关视频。"
    lines = [f"🔍 B站搜索「{keyword}」结果：\n"]
    for i, v in enumerate(results):
        lines.append(
            f"{i+1}. {v['title']}\n"
            f"   up主：{v['author']}  |  {v['play_count']}播放\n"
            f"   {v['link']}\n"
        )
    return "\n".join(lines)


def _bilibili_add_tag_tool(keyword: str) -> str:
    """添加B站兴趣标签工具包装。"""
    if not keyword or not keyword.strip():
        return "请提供要添加的兴趣标签关键词。"
    from utils.bilibili_history import get_bilibili_history
    mgr = get_bilibili_history()
    mgr.add_tag(keyword.strip(), source="ai")
    mgr.save()
    return f"已添加兴趣标签「{keyword.strip()}」，莲心空闲时会去B站搜索相关视频推荐给你~"


def _bilibili_list_tags_tool() -> str:
    """列出所有B站兴趣标签。"""
    from utils.bilibili_history import get_bilibili_history
    mgr = get_bilibili_history()
    active_tags = mgr.get_tags("active")
    paused_tags = mgr.get_tags("paused")
    if not active_tags and not paused_tags:
        return "目前还没有任何B站兴趣标签。你可以说「帮我关注XXX」来添加。"
    lines = ["📋 当前B站兴趣标签：\n"]
    if active_tags:
        lines.append("▸ 活跃标签：")
        for t in active_tags:
            score = t["base_score"] + t.get("boost_score", 0)
            lines.append(f"  · {t['keyword']}（权重 {score}，来源 {t.get('source', 'auto')}）")
    if paused_tags:
        lines.append("▸ 已暂停标签：")
        for t in paused_tags:
            lines.append(f"  · {t['keyword']}（已暂停）")
    return "\n".join(lines)


def _execute_tool_impl(name: str, tool_input: dict) -> str:
    """根据工具名称调用对应的执行函数。调用前检查防御模式。
    支持错误恢复链：网络类工具失败后自动重试+降级。"""
    try:
        # 情感门控必须位于所有工具路由之前，避免 MCP 工具绕过统一检查。
        try:
            from brain.emotional import get_manager as _get_emotion_mgr
            _allowed, _reason = _get_emotion_mgr().check_tool_allowed(name)
            if not _allowed:
                return f"[拒绝] {_reason}"
        except Exception:
            pass
        # ── MCP 工具路由 ──────────────────────────────────────
        if name.startswith("mcp__"):
            try:
                from brain.mcp.mcp_bridge import wrap_as_sync
                return wrap_as_sync(name, tool_input)
            except Exception as e:
                return f"MCP工具调用失败: {e}"

        executor = TOOL_EXECUTORS.get(name)
        if not executor:
            return f"未知工具: {name}"

        # ── 错误恢复链 ──────────────────────────────────────
        from brain.tool_recovery import should_recover, execute_with_recovery

        def _exec(name_to_run, args_to_run):
            ex = TOOL_EXECUTORS.get(name_to_run)
            if not ex:
                return f"未知工具: {name_to_run}"
            try:
                return ex(args_to_run)
            except KeyError as ke:
                return f"参数错误：缺少必需参数 '{ke.args[0]}'，请检查工具定义后重新调用"
            except TypeError as te:
                return f"参数错误：{te}，请检查参数名和类型是否正确"

        try:
            if should_recover(name):
                result, retries, log = execute_with_recovery(name, tool_input, _exec)
                if log:
                    pass
                return result
            else:
                return executor(tool_input)
        except KeyError as ke:
            return f"参数错误：缺少必需参数 '{ke.args[0]}'，请检查工具定义后重新调用"
        except TypeError as te:
            return f"参数错误：{te}，请检查参数名和类型是否正确"
    except Exception:
        raise
_WORKFLOW_CACHE_TTL = {
    "web_search": 15 * 60,
    "fetch_webpage": 60 * 60,
    "fetch_webpage_via_api": 60 * 60,
    "fetch_webpage_browser": 30 * 60,
}


def execute_tool(name: str, tool_input: dict, *, invocation_mode: str = "auto",
                 channel: str = "") -> str:
    """Execute one tool with persistent Workflow step audit and safe read cache."""
    import uuid

    run_id = 0
    context_step_key = ""
    store = None
    step_id = 0
    started = time.perf_counter()
    usage_status = ""
    try:
        from brain.workflow import get_workflow_context, get_workflow_store

        run_id, context_step_key = get_workflow_context()
        if run_id:
            store = get_workflow_store()
            step_id = store.start_step(
                run_id,
                step_key=context_step_key or f"tool:{name}:{uuid.uuid4().hex}",
                name=name,
                kind="tool",
                input_data=tool_input,
            )
            if store.is_cancel_requested(run_id):
                result = "[CANCELLED] 工作流已取消"
                usage_status = "cancelled"
                store.finish_step(step_id, status="cancelled", output_preview=result)
                return result

        ttl = _WORKFLOW_CACHE_TTL.get(name, 0)
        if ttl:
            cache_store = store or get_workflow_store()
            cached_result = cache_store.get_cache(name, tool_input)
            if cached_result is not None:
                usage_status = "cached"
                if step_id:
                    cache_store.finish_step(
                        step_id, status="success", output_preview=cached_result,
                        duration_ms=(time.perf_counter() - started) * 1000, cached=True,
                    )
                return cached_result

        result_text = str(_execute_tool_impl(name, tool_input) or "")
        from brain.tool_usage import classify_tool_result
        usage_status = classify_tool_result(result_text)
        is_error = usage_status not in ("success", "cached")
        if ttl and not is_error:
            (store or get_workflow_store()).put_cache(name, tool_input, result_text, ttl_seconds=ttl)
        if step_id and store:
            store.finish_step(
                step_id, status="failed" if is_error else "success",
                output_preview=result_text, error=result_text if is_error else "",
                duration_ms=(time.perf_counter() - started) * 1000,
            )
            for key in ("path", "file_path", "output_path"):
                raw_path = tool_input.get(key)
                if not raw_path:
                    continue
                try:
                    artifact_path = Path(str(raw_path)).resolve()
                    if artifact_path.is_file():
                        store.add_artifact(
                            run_id, step_id=step_id, artifact_type="file",
                            name=artifact_path.name, uri=str(artifact_path),
                            metadata={"tool": name, "argument": key},
                        )
                except Exception:
                    pass
        return result_text
    except Exception as exc:
        usage_status = "failure"
        if step_id and store:
            store.finish_step(
                step_id, status="failed", error=str(exc),
                duration_ms=(time.perf_counter() - started) * 1000,
            )
        raise
    finally:
        if usage_status:
            try:
                from brain.tool_usage import get_tool_usage_store

                core_names = {
                    item.get("function", {}).get("name", "") for item in TOOL_DEFINITIONS
                }
                if name.startswith("mcp__"):
                    source_kind = "mcp"
                    provider_id = name.split("__", 2)[1]
                elif name in core_names:
                    source_kind = "builtin"
                    provider_id = "lianxin"
                else:
                    source_kind = "skill"
                    provider_id = "skill"
                    try:
                        from brain.skill_manager import _skill_registry
                        for skill_name, info in _skill_registry.items():
                            names = {
                                item.get("function", {}).get("name", "")
                                for item in info.get("tool_definitions", [])
                            }
                            if name in names:
                                provider_id = skill_name
                                break
                    except Exception:
                        pass
                context = getattr(_tool_context, "cross_session", {}) or {}
                get_tool_usage_store().record(
                    name, status=usage_status,
                    duration_ms=(time.perf_counter() - started) * 1000,
                    source_kind=source_kind, provider_id=provider_id,
                    invocation_mode=invocation_mode,
                    session_id=context.get("session_id"), workflow_run_id=run_id or None,
                    channel=channel or context.get("source_channel", ""),
                )
            except Exception:
                pass

# Read-only runtime capability lookup.  CapabilityCenter reads the same catalog.
CAPABILITY_QUERY_DEFINITION = {
    "type": "function",
    "function": {
        "name": "query_capabilities",
        "description": (
            "仅当用户明确询问莲心会什么、有哪些功能、能否支持某项操作时调用。"
            "查询当前实时能力目录，区分可直接使用、已停用和当前不可用。"
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "想了解的能力或操作；可为空。"},
                "category": {"type": "string", "description": "可选能力分类。"},
                "limit": {"type": "integer", "description": "返回条数，默认 20，最大 50。"},
            },
            "required": [],
        },
    },
}
TOOL_DEFINITIONS.append(CAPABILITY_QUERY_DEFINITION)
TOOL_EXECUTORS["query_capabilities"] = lambda inp: __import__(
    "brain.capability_knowledge", fromlist=["format_capability_query"]
).format_capability_query(inp.get("query", ""), inp.get("category", ""), inp.get("limit", 20))

# ── 初始化工具注册中心（模块导入时自动注册所有工具）─────────
try:
    from brain.tool_registry import init_tool_registry
    init_tool_registry(list(TOOL_EXECUTORS.keys()))
except Exception:
    pass